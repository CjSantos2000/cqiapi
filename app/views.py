import tempfile
import time

from django.shortcuts import render
from docx import Document
from django.http import FileResponse
import os
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import pandas as pd
from django.http import HttpResponse
from django.core.files.base import ContentFile
from io import BytesIO
from .models import *
import requests
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from google.oauth2 import service_account
from gspread_formatting import *


def replace_text_v1(doc, replacement_dict):
    def replace_in_paragraph(paragraph, replacement_dict):
        for run in paragraph.runs:
            for old_text, new_text in replacement_dict.items():
                if f"##{old_text}" in run.text:
                    run.text = run.text.replace(f"##{old_text}", new_text)
                    print(f"replace_in_paragraph {run.text}")

    def replace_in_table(table, replacement_dict):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacement_dict)
                    print(f"replace_in_table {paragraph.text}")

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacement_dict)

    for table in doc.tables:
        replace_in_table(table, replacement_dict)


def save_document_with_temp(doc, local_path):
    # Ensure the directory exists
    directory = os.path.dirname(local_path)
    if not os.path.exists(directory):
        os.makedirs(directory)
        print(f"Directory {directory} created")

    # check if local_path file exists
    if os.path.exists(local_path):
        os.remove(local_path)
        print(f"File {local_path} removed")

    # Log directory permissions
    print(f"Directory permissions for {directory}: {os.access(directory, os.W_OK)}")

    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_file_path = tmp_file.name

    # Save the document to the temporary file
    doc.save(tmp_file_path)

    try:
        # Move the temporary file to the desired local path
        shutil.move(tmp_file_path, local_path)
        print(f"Document saved to {local_path}")
    except PermissionError as e:
        print(f"Failed to save the document: {e}")
    finally:
        # Ensure the temporary file is removed
        if os.path.exists(tmp_file_path):
            os.remove(tmp_file_path)


@csrf_exempt
def download_obe(request):
    if request.method == "POST":
        data = json.loads(request.body)

        user_id = data["user_id"]
        form_name = "obe"
        fileid = data["fileid"]
        script_directory = os.path.dirname(os.path.abspath(__file__))
        output_file_path = os.path.join(
            script_directory, f"obe-{user_id}-{fileid}.docx"
        )

        print("script_directory", script_directory)

        response = FileResponse(open(output_file_path, "rb"))

        print(response)

        record = DownloadModel.objects.filter(
            user_id=user_id, record_id=fileid, form_name=form_name
        ).count()
        if record > 0:
            script_directory = os.path.dirname(os.path.abspath(__file__))
            output_file_path = os.path.join(
                script_directory, f"obe-{user_id}-{fileid}.docx"
            )

            response = FileResponse(open(output_file_path, "rb"))
            response["Content-Disposition"] = 'attachment; filename="obe-output.docx"'
            return response
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def download_matrix(request):
    if request.method == "POST":
        data = json.loads(request.body)

        user_id = data["user_id"]
        form_name = "matrix"
        fileid = data["fileid"]

        record = DownloadModel.objects.filter(
            user_id=user_id, record_id=fileid, form_name=form_name
        ).count()
        if record > 0:
            script_directory = os.path.dirname(os.path.abspath(__file__))
            output_file_path = os.path.join(
                script_directory, f"cam-{user_id}-{fileid}.docx"
            )

            response = FileResponse(open(output_file_path, "rb"))
            response["Content-Disposition"] = 'attachment; filename="obe-output.docx"'
            return response
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def download_datasheet(request):
    if request.method == "POST":
        data = json.loads(request.body)

        user_id = data["user_id"]
        form_name = "datasheet"
        fileid = data["fileid"]

        record = DownloadModel.objects.filter(
            user_id=user_id, record_id=fileid, form_name=form_name
        ).count()
        if record > 0:
            script_directory = os.path.dirname(os.path.abspath(__file__))
            output_file_path = os.path.join(
                script_directory, f"datasheet-{user_id}-{fileid}.docx"
            )

            response = FileResponse(open(output_file_path, "rb"))
            response["Content-Disposition"] = 'attachment; filename="obe-output.docx"'
            return response
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def download_summary(request):
    if request.method == "POST":
        data = json.loads(request.body)

        user_id = data["user_id"]
        form_name = "summary"
        fileid = data["fileid"]

        record = DownloadModel.objects.filter(
            user_id=user_id, record_id=fileid, form_name=form_name
        ).count()
        if record > 0:
            script_directory = os.path.dirname(os.path.abspath(__file__))
            output_file_path = os.path.join(
                script_directory, f"cas-{user_id}-{fileid}.docx"
            )

            response = FileResponse(open(output_file_path, "rb"))
            response["Content-Disposition"] = 'attachment; filename="obe-output.docx"'
            return response
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def download_tos(request):
    if request.method == "POST":
        data = json.loads(request.body)

        user_id = data["user_id"]
        form_name = "tos"
        fileid = data["fileid"]

        record = DownloadModel.objects.filter(
            user_id=user_id, record_id=fileid, form_name=form_name
        ).count()
        if record > 0:
            script_directory = os.path.dirname(os.path.abspath(__file__))
            output_file_path = os.path.join(
                script_directory, f"tos-{user_id}-{fileid}.docx"
            )

            response = FileResponse(open(output_file_path, "rb"))
            response["Content-Disposition"] = 'attachment; filename="obe-output.docx"'
            return response
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def download_plo(request):
    if request.method == "POST":
        data = json.loads(request.body)

        user_id = data["user_id"]
        form_name = "plo"
        fileid = data["fileid"]

        record = DownloadModel.objects.filter(
            user_id=user_id, record_id=fileid, form_name=form_name
        ).count()
        if record > 0:
            script_directory = os.path.dirname(os.path.abspath(__file__))
            output_file_path = os.path.join(
                script_directory, f"plo-{user_id}-{fileid}.docx"
            )

            response = FileResponse(open(output_file_path, "rb"))
            response["Content-Disposition"] = 'attachment; filename="obe-output.docx"'
            return response
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def download_record(request):
    if request.method == "POST":
        data = json.loads(request.body)

        user_id = data["user_id"]
        form_name = "class-record"
        fileid = data["fileid"]

        record = DownloadModel.objects.filter(
            user_id=user_id, record_id=fileid, form_name=form_name
        ).count()
        if record > 0:
            script_directory = os.path.dirname(os.path.abspath(__file__))
            output_file_path = os.path.join(
                script_directory, f"Class-record-{user_id}-{fileid}.xlsx"
            )

            response = FileResponse(open(output_file_path, "rb"))
            response["Content-Disposition"] = 'attachment; filename="Class-record.xlsx"'
            return response
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def obe(request):
    if request.method == "POST":
        start = time.time()
        print("Start-Time: ", start)
        data = json.loads(request.body)

        class_record = data["class_record_auto_fetch_data"]
        lastid = data["lastid"]
        user_id = data["user_id"]
        obj, created = AutoFetchModel.objects.update_or_create(
            user_id=user_id, category="auto-fetch-data", defaults={"name": class_record}
        )

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"course_title": "a"},
            {"course_code": "a"},
            {"pre_requisite": "a"},
            {"co_requisite": "a"},
            {"number_units": "a"},
            {"course_description": ""},
            {"course_outcomes": ""},
            {"course_objectives": ""},
            {"po_a": ""},
            {"a_peo1": ""},
            {"a_peo2": ""},
            {"a_peo3": ""},
            {"po_b": ""},
            {"b_peo1": ""},
            {"b_peo2": ""},
            {"b_peo3": ""},
            {"po_c": ""},
            {"c_peo1": ""},
            {"c_peo2": ""},
            {"c_peo3": ""},
            {"po_d": ""},
            {"d_peo1": ""},
            {"d_peo2": ""},
            {"d_peo3": ""},
            {"po_e": ""},
            {"e_peo1": ""},
            {"e_peo2": ""},
            {"e_peo3": ""},
            {"po_f": ""},
            {"f_peo1": ""},
            {"f_peo2": ""},
            {"f_peo3": ""},
            {"po_g": ""},
            {"g_peo1": ""},
            {"g_peo2": ""},
            {"g_peo3": ""},
            {"po_h": ""},
            {"h_peo1": ""},
            {"h_peo2": ""},
            {"h_peo3": ""},
            {"po_i": ""},
            {"i_peo1": ""},
            {"i_peo2": ""},
            {"i_peo3": ""},
            {"po_j": ""},
            {"j_peo1": ""},
            {"j_peo2": ""},
            {"j_peo3": ""},
            {"po_k": ""},
            {"k_peo1": ""},
            {"k_peo2": ""},
            {"k_peo3": ""},
            {"po_l": ""},
            {"l_peo1": ""},
            {"l_peo2": ""},
            {"l_peo3": ""},
            {"po_m": ""},
            {"m_peo1": ""},
            {"m_peo2": ""},
            {"m_peo3": ""},
            {"clo-1": ""},
            {"a-1": ""},
            {"b-1": ""},
            {"c-1": ""},
            {"d-1": ""},
            {"e-1": ""},
            {"f-1": ""},
            {"g-1": ""},
            {"h-1": ""},
            {"i-1": ""},
            {"j-1": ""},
            {"k-1": ""},
            {"l-1": ""},
            {"m-1": ""},
            {"clo-2": ""},
            {"a-2": ""},
            {"b-2": ""},
            {"c-2": ""},
            {"d-2": ""},
            {"e-2": ""},
            {"f-2": ""},
            {"g-2": ""},
            {"h-2": ""},
            {"i-2": ""},
            {"j-2": ""},
            {"k-2": ""},
            {"l-2": ""},
            {"m-2": ""},
            {"clo-3": ""},
            {"a-3": ""},
            {"b-3": ""},
            {"c-3": ""},
            {"d-3": ""},
            {"e-3": ""},
            {"f-3": ""},
            {"g-3": ""},
            {"h-3": ""},
            {"i-3": ""},
            {"j-3": ""},
            {"k-3": ""},
            {"l-3": ""},
            {"m-3": ""},
            {"clo-4": ""},
            {"a-4": ""},
            {"b-4": ""},
            {"c-4": ""},
            {"d-4": ""},
            {"e-4": ""},
            {"f-4": ""},
            {"g-4": ""},
            {"h-4": ""},
            {"i-4": ""},
            {"j-4": ""},
            {"k-4": ""},
            {"l-4": ""},
            {"m-4": ""},
            {"clo-5": ""},
            {"a-5": ""},
            {"b-5": ""},
            {"c-5": ""},
            {"d-5": ""},
            {"e-5": ""},
            {"f-5": ""},
            {"g-5": ""},
            {"h-5": ""},
            {"i-5": ""},
            {"j-5": ""},
            {"k-5": ""},
            {"l-5": ""},
            {"m-5": ""},
            {"clo-6": ""},
            {"a-6": ""},
            {"b-6": ""},
            {"c-6": ""},
            {"d-6": ""},
            {"e-6": ""},
            {"f-6": ""},
            {"g-6": ""},
            {"h-6": ""},
            {"i-6": ""},
            {"j-6": ""},
            {"k-6": ""},
            {"l-6": ""},
            {"m-6": ""},
            {"clo-7": ""},
            {"a-7": ""},
            {"b-7": ""},
            {"c-7": ""},
            {"d-7": ""},
            {"e-7": ""},
            {"f-7": ""},
            {"g-7": ""},
            {"h-7": ""},
            {"i-7": ""},
            {"j-7": ""},
            {"k-7": ""},
            {"l-7": ""},
            {"m-7": ""},
            {"clo-8": ""},
            {"a-8": ""},
            {"b-8": ""},
            {"c-8": ""},
            {"d-8": ""},
            {"e-8": ""},
            {"f-8": ""},
            {"g-8": ""},
            {"h-8": ""},
            {"i-8": ""},
            {"j-8": ""},
            {"k-8": ""},
            {"l-8": ""},
            {"m-8": ""},
            {"clo-9": ""},
            {"a-9": ""},
            {"b-9": ""},
            {"c-9": ""},
            {"d-9": ""},
            {"e-9": ""},
            {"f-9": ""},
            {"g-9": ""},
            {"h-9": ""},
            {"i-9": ""},
            {"j-9": ""},
            {"k-9": ""},
            {"l-9": ""},
            {"m-9": ""},
            {"clo-10": ""},
            {"a-10": ""},
            {"b-10": ""},
            {"c-10": ""},
            {"d-10": ""},
            {"e-10": ""},
            {"f-10": ""},
            {"g-10": ""},
            {"h-10": ""},
            {"i-10": ""},
            {"j-10": ""},
            {"k-10": ""},
            {"l-10": ""},
            {"m-10": ""},
            {"plo_i": ""},
            {"peo1_i": ""},
            {"peo2_i": ""},
            {"peo3_i": ""},
            {"plo_e": ""},
            {"peo1_e": ""},
            {"peo2_e": ""},
            {"peo3_e": ""},
            {"plo_d": ""},
            {"peo1_d": ""},
            {"peo2_d": ""},
            {"peo3_d": ""},
            {"pea1_mission_a": ""},
            {"pea1_mission_b": ""},
            {"pea1_mission_c": ""},
            {"pea1_mission_d": ""},
            {"pea2_mission_a": ""},
            {"pea2_mission_b": ""},
            {"pea2_mission_c": ""},
            {"pea2_mission_d": ""},
            {"pea3_mission_a": ""},
            {"pea3_mission_b": ""},
            {"pea3_mission_c": ""},
            {"pea3_mission_d": ""},
            {"clo-obj-2-1": ""},
            {"prelim-exam-1": ""},
            {"exams1-1": ""},
            {"projects1-1": ""},
            {"presentations1-1": ""},
            {"assignments1-1": ""},
            {"midterm-exam-1": ""},
            {"exams2-1": ""},
            {"projects2-1": ""},
            {"presentations2-1": ""},
            {"assignments2-1": ""},
            {"final-exam-1": ""},
            {"exams3-1": ""},
            {"projects3-1": ""},
            {"presentations3-1": ""},
            {"assignments3-1": ""},
            {"compre-exam-1": ""},
            {"exams4-1": ""},
            {"projects4-1": ""},
            {"presentations4-1": ""},
            {"assignments4-1": ""},
            {"clo-obj-2-2": ""},
            {"prelim-exam-2": ""},
            {"exams1-2": ""},
            {"projects1-2": ""},
            {"presentations1-2": ""},
            {"assignments1-2": ""},
            {"midterm-exam-2": ""},
            {"exams2-2": ""},
            {"projects2-2": ""},
            {"presentations2-2": ""},
            {"assignments2-2": ""},
            {"final-exam-2": ""},
            {"exams3-2": ""},
            {"projects3-2": ""},
            {"presentations3-2": ""},
            {"assignments3-2": ""},
            {"compre-exam-2": ""},
            {"exams4-2": ""},
            {"projects4-2": ""},
            {"presentations4-2": ""},
            {"assignments4-2": ""},
            {"clo-obj-2-3": ""},
            {"prelim-exam-3": ""},
            {"exams1-3": ""},
            {"projects1-3": ""},
            {"presentations1-3": ""},
            {"assignments1-3": ""},
            {"midterm-exam-3": ""},
            {"exams2-3": ""},
            {"projects2-3": ""},
            {"presentations2-3": ""},
            {"assignments2-3": ""},
            {"final-exam-3": ""},
            {"exams3-3": ""},
            {"projects3-3": ""},
            {"presentations3-3": ""},
            {"assignments3-3": ""},
            {"compre-exam-3": ""},
            {"exams4-3": ""},
            {"projects4-3": ""},
            {"presentations4-3": ""},
            {"assignments4-3": ""},
            {"date-week-1-1st": ""},
            {"date-week-1-2nd": ""},
            {"course-1-obe": ""},
            {"cilo-1-obe": ""},
            {"topics-1-obe": ""},
            {"tla-1-obe": ""},
            {"tasks-1-obe": ""},
            {"date-week-2-1st": ""},
            {"date-week-2-2nd": ""},
            {"course-2-obe": ""},
            {"cilo-2-obe": ""},
            {"topics-2-obe": ""},
            {"tla-2-obe": ""},
            {"tasks-2-obe": ""},
            {"date-week-3-1st": ""},
            {"date-week-3-2nd": ""},
            {"course-3-obe": ""},
            {"cilo-3-obe": ""},
            {"topics-3-obe": ""},
            {"tla-3-obe": ""},
            {"tasks-3-obe": ""},
            {"lrass": ""},
            {"date-week-4-1st": ""},
            {"date-week-4-2nd": ""},
            {"date-week-5-1st": ""},
            {"date-week-5-2nd": ""},
            {"date-week-6-1st": ""},
            {"date-week-6-2nd": ""},
            {"date-week-7-1st": ""},
            {"date-week-7-2nd": ""},
            {"date-week-8-1st": ""},
            {"date-week-8-2nd": ""},
            {"date-week-9-1st": ""},
            {"date-week-9-2nd": ""},
            {"cilo-4-obe": ""},
            {"cilo-5-obe": ""},
            {"cilo-6-obe": ""},
            {"cilo-7-obe": ""},
            {"cilo-8-obe": ""},
            {"cilo-9-obe": ""},
            {"cilo-10-obe": ""},
            {"cilo-11-obe": ""},
            {"cilo-12-obe": ""},
            {"cilo-13-obe": ""},
            {"cilo-14-obe": ""},
            {"cilo-15-obe": ""},
            {"cilo-16-obe": ""},
            {"cilo-17-obe": ""},
            {"cilo-18-obe": ""},
            {"course-4-obe": ""},
            {"course-5-obe": ""},
            {"course-6-obe": ""},
            {"course-7-obe": ""},
            {"course-8-obe": ""},
            {"course-9-obe": ""},
            {"course-10-obe": ""},
            {"course-11-obe": ""},
            {"course-12-obe": ""},
            {"course-13-obe": ""},
            {"course-14-obe": ""},
            {"course-15-obe": ""},
            {"course-16-obe": ""},
            {"course-17-obe": ""},
            {"course-18-obe": ""},
            {"tla-4-obe": ""},
            {"tla-5-obe": ""},
            {"tla-6-obe": ""},
            {"tla-7-obe": ""},
            {"tla-8-obe": ""},
            {"tla-9-obe": ""},
            {"tla-10-obe": ""},
            {"tla-11-obe": ""},
            {"tla-12-obe": ""},
            {"tla-13-obe": ""},
            {"tla-14-obe": ""},
            {"tla-15-obe": ""},
            {"tla-16-obe": ""},
            {"tla-17-obe": ""},
            {"tla-18-obe": ""},
            {"tasks-4-obe": ""},
            {"tasks-5-obe": ""},
            {"tasks-6-obe": ""},
            {"tasks-7-obe": ""},
            {"tasks-8-obe": ""},
            {"tasks-9-obe": ""},
            {"tasks-10-obe": ""},
            {"tasks-11-obe": ""},
            {"tasks-12-obe": ""},
            {"tasks-13-obe": ""},
            {"tasks-14-obe": ""},
            {"tasks-15-obe": ""},
            {"tasks-16-obe": ""},
            {"tasks-17-obe": ""},
            {"tasks-18-obe": ""},
            {"topics-4-obe": ""},
            {"topics-5-obe": ""},
            {"topics-6-obe": ""},
            {"topics-7-obe": ""},
            {"topics-8-obe": ""},
            {"topics-9-obe": ""},
            {"topics-10-obe": ""},
            {"topics-11-obe": ""},
            {"topics-12-obe": ""},
            {"topics-13-obe": ""},
            {"topics-14-obe": ""},
            {"topics-15-obe": ""},
            {"topics-16-obe": ""},
            {"topics-17-obe": ""},
            {"topics-18-obe": ""},
            {"date-week-10-1st": ""},
            {"date-week-11-1st": ""},
            {"date-week-12-1st": ""},
            {"date-week-13-1st": ""},
            {"date-week-14-1st": ""},
            {"date-week-15-1st": ""},
            {"date-week-16-1st": ""},
            {"date-week-17-1st": ""},
            {"date-week-18-1st": ""},
            {"date-week-10-2nd": ""},
            {"date-week-11-2nd": ""},
            {"date-week-12-2nd": ""},
            {"date-week-13-2nd": ""},
            {"date-week-14-2nd": ""},
            {"date-week-15-2nd": ""},
            {"date-week-16-2nd": ""},
            {"date-week-17-2nd": ""},
            {"date-week-18-2nd": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "obe.docx")
        output_file_path = os.path.join(
            script_directory, f"obe-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        a = DownloadModel(user_id=user_id, record_id=lastid, form_name="obe")
        a.save()

        end = time.time()
        print("End-Time: ", end)
        print("Time: ", end - start)

        response = FileResponse(open(output_file_path, "rb"))
        response["Content-Disposition"] = 'attachment; filename="datasheet-output.docx"'
        return response
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def obe_update(request):
    if request.method == "POST":
        data = json.loads(request.body)

        class_record = data["class_record_auto_fetch_data"]
        lastid = data["updateid"]
        user_id = data["user_id"]

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"course_title": "a"},
            {"course_code": "a"},
            {"pre_requisite": "a"},
            {"co_requisite": "a"},
            {"number_units": "a"},
            {"course_description": ""},
            {"course_outcomes": ""},
            {"course_objectives": ""},
            {"po_a": ""},
            {"a_peo1": ""},
            {"a_peo2": ""},
            {"a_peo3": ""},
            {"po_b": ""},
            {"b_peo1": ""},
            {"b_peo2": ""},
            {"b_peo3": ""},
            {"po_c": ""},
            {"c_peo1": ""},
            {"c_peo2": ""},
            {"c_peo3": ""},
            {"po_d": ""},
            {"d_peo1": ""},
            {"d_peo2": ""},
            {"d_peo3": ""},
            {"po_e": ""},
            {"e_peo1": ""},
            {"e_peo2": ""},
            {"e_peo3": ""},
            {"po_f": ""},
            {"f_peo1": ""},
            {"f_peo2": ""},
            {"f_peo3": ""},
            {"po_g": ""},
            {"g_peo1": ""},
            {"g_peo2": ""},
            {"g_peo3": ""},
            {"po_h": ""},
            {"h_peo1": ""},
            {"h_peo2": ""},
            {"h_peo3": ""},
            {"po_i": ""},
            {"i_peo1": ""},
            {"i_peo2": ""},
            {"i_peo3": ""},
            {"po_j": ""},
            {"j_peo1": ""},
            {"j_peo2": ""},
            {"j_peo3": ""},
            {"po_k": ""},
            {"k_peo1": ""},
            {"k_peo2": ""},
            {"k_peo3": ""},
            {"po_l": ""},
            {"l_peo1": ""},
            {"l_peo2": ""},
            {"l_peo3": ""},
            {"po_m": ""},
            {"m_peo1": ""},
            {"m_peo2": ""},
            {"m_peo3": ""},
            {"clo-1": ""},
            {"a-1": ""},
            {"b-1": ""},
            {"c-1": ""},
            {"d-1": ""},
            {"e-1": ""},
            {"f-1": ""},
            {"g-1": ""},
            {"h-1": ""},
            {"i-1": ""},
            {"j-1": ""},
            {"k-1": ""},
            {"l-1": ""},
            {"m-1": ""},
            {"clo-2": ""},
            {"a-2": ""},
            {"b-2": ""},
            {"c-2": ""},
            {"d-2": ""},
            {"e-2": ""},
            {"f-2": ""},
            {"g-2": ""},
            {"h-2": ""},
            {"i-2": ""},
            {"j-2": ""},
            {"k-2": ""},
            {"l-2": ""},
            {"m-2": ""},
            {"clo-3": ""},
            {"a-3": ""},
            {"b-3": ""},
            {"c-3": ""},
            {"d-3": ""},
            {"e-3": ""},
            {"f-3": ""},
            {"g-3": ""},
            {"h-3": ""},
            {"i-3": ""},
            {"j-3": ""},
            {"k-3": ""},
            {"l-3": ""},
            {"m-3": ""},
            {"clo-4": ""},
            {"a-4": ""},
            {"b-4": ""},
            {"c-4": ""},
            {"d-4": ""},
            {"e-4": ""},
            {"f-4": ""},
            {"g-4": ""},
            {"h-4": ""},
            {"i-4": ""},
            {"j-4": ""},
            {"k-4": ""},
            {"l-4": ""},
            {"m-4": ""},
            {"clo-5": ""},
            {"a-5": ""},
            {"b-5": ""},
            {"c-5": ""},
            {"d-5": ""},
            {"e-5": ""},
            {"f-5": ""},
            {"g-5": ""},
            {"h-5": ""},
            {"i-5": ""},
            {"j-5": ""},
            {"k-5": ""},
            {"l-5": ""},
            {"m-5": ""},
            {"clo-6": ""},
            {"a-6": ""},
            {"b-6": ""},
            {"c-6": ""},
            {"d-6": ""},
            {"e-6": ""},
            {"f-6": ""},
            {"g-6": ""},
            {"h-6": ""},
            {"i-6": ""},
            {"j-6": ""},
            {"k-6": ""},
            {"l-6": ""},
            {"m-6": ""},
            {"clo-7": ""},
            {"a-7": ""},
            {"b-7": ""},
            {"c-7": ""},
            {"d-7": ""},
            {"e-7": ""},
            {"f-7": ""},
            {"g-7": ""},
            {"h-7": ""},
            {"i-7": ""},
            {"j-7": ""},
            {"k-7": ""},
            {"l-7": ""},
            {"m-7": ""},
            {"clo-8": ""},
            {"a-8": ""},
            {"b-8": ""},
            {"c-8": ""},
            {"d-8": ""},
            {"e-8": ""},
            {"f-8": ""},
            {"g-8": ""},
            {"h-8": ""},
            {"i-8": ""},
            {"j-8": ""},
            {"k-8": ""},
            {"l-8": ""},
            {"m-8": ""},
            {"clo-9": ""},
            {"a-9": ""},
            {"b-9": ""},
            {"c-9": ""},
            {"d-9": ""},
            {"e-9": ""},
            {"f-9": ""},
            {"g-9": ""},
            {"h-9": ""},
            {"i-9": ""},
            {"j-9": ""},
            {"k-9": ""},
            {"l-9": ""},
            {"m-9": ""},
            {"clo-10": ""},
            {"a-10": ""},
            {"b-10": ""},
            {"c-10": ""},
            {"d-10": ""},
            {"e-10": ""},
            {"f-10": ""},
            {"g-10": ""},
            {"h-10": ""},
            {"i-10": ""},
            {"j-10": ""},
            {"k-10": ""},
            {"l-10": ""},
            {"m-10": ""},
            {"plo_i": ""},
            {"peo1_i": ""},
            {"peo2_i": ""},
            {"peo3_i": ""},
            {"plo_e": ""},
            {"peo1_e": ""},
            {"peo2_e": ""},
            {"peo3_e": ""},
            {"plo_d": ""},
            {"peo1_d": ""},
            {"peo2_d": ""},
            {"peo3_d": ""},
            {"pea1_mission_a": ""},
            {"pea1_mission_b": ""},
            {"pea1_mission_c": ""},
            {"pea1_mission_d": ""},
            {"pea2_mission_a": ""},
            {"pea2_mission_b": ""},
            {"pea2_mission_c": ""},
            {"pea2_mission_d": ""},
            {"pea3_mission_a": ""},
            {"pea3_mission_b": ""},
            {"pea3_mission_c": ""},
            {"pea3_mission_d": ""},
            {"clo-obj-2-1": ""},
            {"prelim-exam-1": ""},
            {"exams1-1": ""},
            {"projects1-1": ""},
            {"presentations1-1": ""},
            {"assignments1-1": ""},
            {"midterm-exam-1": ""},
            {"exams2-1": ""},
            {"projects2-1": ""},
            {"presentations2-1": ""},
            {"assignments2-1": ""},
            {"final-exam-1": ""},
            {"exams3-1": ""},
            {"projects3-1": ""},
            {"presentations3-1": ""},
            {"assignments3-1": ""},
            {"compre-exam-1": ""},
            {"exams4-1": ""},
            {"projects4-1": ""},
            {"presentations4-1": ""},
            {"assignments4-1": ""},
            {"clo-obj-2-2": ""},
            {"prelim-exam-2": ""},
            {"exams1-2": ""},
            {"projects1-2": ""},
            {"presentations1-2": ""},
            {"assignments1-2": ""},
            {"midterm-exam-2": ""},
            {"exams2-2": ""},
            {"projects2-2": ""},
            {"presentations2-2": ""},
            {"assignments2-2": ""},
            {"final-exam-2": ""},
            {"exams3-2": ""},
            {"projects3-2": ""},
            {"presentations3-2": ""},
            {"assignments3-2": ""},
            {"compre-exam-2": ""},
            {"exams4-2": ""},
            {"projects4-2": ""},
            {"presentations4-2": ""},
            {"assignments4-2": ""},
            {"clo-obj-2-3": ""},
            {"prelim-exam-3": ""},
            {"exams1-3": ""},
            {"projects1-3": ""},
            {"presentations1-3": ""},
            {"assignments1-3": ""},
            {"midterm-exam-3": ""},
            {"exams2-3": ""},
            {"projects2-3": ""},
            {"presentations2-3": ""},
            {"assignments2-3": ""},
            {"final-exam-3": ""},
            {"exams3-3": ""},
            {"projects3-3": ""},
            {"presentations3-3": ""},
            {"assignments3-3": ""},
            {"compre-exam-3": ""},
            {"exams4-3": ""},
            {"projects4-3": ""},
            {"presentations4-3": ""},
            {"assignments4-3": ""},
            {"date-week-1-1st": ""},
            {"date-week-1-2nd": ""},
            {"course-1-obe": ""},
            {"cilo-1-obe": ""},
            {"topics-1-obe": ""},
            {"tla-1-obe": ""},
            {"tasks-1-obe": ""},
            {"date-week-2-1st": ""},
            {"date-week-2-2nd": ""},
            {"course-2-obe": ""},
            {"cilo-2-obe": ""},
            {"topics-2-obe": ""},
            {"tla-2-obe": ""},
            {"tasks-2-obe": ""},
            {"date-week-3-1st": ""},
            {"date-week-3-2nd": ""},
            {"course-3-obe": ""},
            {"cilo-3-obe": ""},
            {"topics-3-obe": ""},
            {"tla-3-obe": ""},
            {"tasks-3-obe": ""},
            {"lrass": ""},
            {"date-week-4-1st": ""},
            {"date-week-4-2nd": ""},
            {"date-week-5-1st": ""},
            {"date-week-5-2nd": ""},
            {"date-week-6-1st": ""},
            {"date-week-6-2nd": ""},
            {"date-week-7-1st": ""},
            {"date-week-7-2nd": ""},
            {"date-week-8-1st": ""},
            {"date-week-8-2nd": ""},
            {"date-week-9-1st": ""},
            {"date-week-9-2nd": ""},
            {"cilo-4-obe": ""},
            {"cilo-5-obe": ""},
            {"cilo-6-obe": ""},
            {"cilo-7-obe": ""},
            {"cilo-8-obe": ""},
            {"cilo-9-obe": ""},
            {"cilo-10-obe": ""},
            {"cilo-11-obe": ""},
            {"cilo-12-obe": ""},
            {"cilo-13-obe": ""},
            {"cilo-14-obe": ""},
            {"cilo-15-obe": ""},
            {"cilo-16-obe": ""},
            {"cilo-17-obe": ""},
            {"cilo-18-obe": ""},
            {"course-4-obe": ""},
            {"course-5-obe": ""},
            {"course-6-obe": ""},
            {"course-7-obe": ""},
            {"course-8-obe": ""},
            {"course-9-obe": ""},
            {"course-10-obe": ""},
            {"course-11-obe": ""},
            {"course-12-obe": ""},
            {"course-13-obe": ""},
            {"course-14-obe": ""},
            {"course-15-obe": ""},
            {"course-16-obe": ""},
            {"course-17-obe": ""},
            {"course-18-obe": ""},
            {"tla-4-obe": ""},
            {"tla-5-obe": ""},
            {"tla-6-obe": ""},
            {"tla-7-obe": ""},
            {"tla-8-obe": ""},
            {"tla-9-obe": ""},
            {"tla-10-obe": ""},
            {"tla-11-obe": ""},
            {"tla-12-obe": ""},
            {"tla-13-obe": ""},
            {"tla-14-obe": ""},
            {"tla-15-obe": ""},
            {"tla-16-obe": ""},
            {"tla-17-obe": ""},
            {"tla-18-obe": ""},
            {"tasks-4-obe": ""},
            {"tasks-5-obe": ""},
            {"tasks-6-obe": ""},
            {"tasks-7-obe": ""},
            {"tasks-8-obe": ""},
            {"tasks-9-obe": ""},
            {"tasks-10-obe": ""},
            {"tasks-11-obe": ""},
            {"tasks-12-obe": ""},
            {"tasks-13-obe": ""},
            {"tasks-14-obe": ""},
            {"tasks-15-obe": ""},
            {"tasks-16-obe": ""},
            {"tasks-17-obe": ""},
            {"tasks-18-obe": ""},
            {"topics-4-obe": ""},
            {"topics-5-obe": ""},
            {"topics-6-obe": ""},
            {"topics-7-obe": ""},
            {"topics-8-obe": ""},
            {"topics-9-obe": ""},
            {"topics-10-obe": ""},
            {"topics-11-obe": ""},
            {"topics-12-obe": ""},
            {"topics-13-obe": ""},
            {"topics-14-obe": ""},
            {"topics-15-obe": ""},
            {"topics-16-obe": ""},
            {"topics-17-obe": ""},
            {"topics-18-obe": ""},
            {"date-week-10-1st": ""},
            {"date-week-11-1st": ""},
            {"date-week-12-1st": ""},
            {"date-week-13-1st": ""},
            {"date-week-14-1st": ""},
            {"date-week-15-1st": ""},
            {"date-week-16-1st": ""},
            {"date-week-17-1st": ""},
            {"date-week-18-1st": ""},
            {"date-week-10-2nd": ""},
            {"date-week-11-2nd": ""},
            {"date-week-12-2nd": ""},
            {"date-week-13-2nd": ""},
            {"date-week-14-2nd": ""},
            {"date-week-15-2nd": ""},
            {"date-week-16-2nd": ""},
            {"date-week-17-2nd": ""},
            {"date-week-18-2nd": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "obe.docx")
        output_file_path = os.path.join(
            script_directory, f"obe-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        try:
            DownloadModel.objects.update_or_create(
                user_id=user_id, record_id=lastid, form_name="obe", defaults={}
            )
        except DownloadModel.MultipleObjectsReturned:
            DownloadModel.objects.filter(
                user_id=user_id, record_id=lastid, form_name="obe"
            ).delete()
            DownloadModel.objects.get_or_create(
                user_id=user_id, record_id=lastid, form_name="obe"
            )

        return JsonResponse({"status": "success", "message": "File created"})

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def cam(request):
    if request.method == "POST":
        start = time.time()
        print("Start-Time: ", start)
        data = json.loads(request.body)

        class_record = data["class_record_auto_fetch_data"]
        user_id = data["user_id"]
        lastid = data["lastid"]

        obj, created = AutoFetchModel.objects.update_or_create(
            user_id=user_id, category="auto-fetch-data", defaults={"name": class_record}
        )

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"course_title": "a"},
            {"course_code": "a"},
            {"pre_requisite": "a"},
            {"prior_courses_with_similar_plo": "a"},
            {"academic_year": "a"},
            {"faculty": "a"},
            {"clo-1": "sdfsd"},
            {"1a": ""},
            {"1b": ""},
            {"1c": ""},
            {"1d": ""},
            {"1e": ""},
            {"1f": ""},
            {"1g": ""},
            {"1h": ""},
            {"1i": ""},
            {"1j": ""},
            {"1k": ""},
            {"1l": ""},
            {"1m": ""},
            {"clo-2": ""},
            {"2a": ""},
            {"2b": ""},
            {"2c": ""},
            {"2d": ""},
            {"2e": ""},
            {"2f": ""},
            {"2g": ""},
            {"2h": ""},
            {"2i": ""},
            {"2j": ""},
            {"2k": ""},
            {"2l": ""},
            {"2m": ""},
            {"clo-3": ""},
            {"3a": ""},
            {"3b": ""},
            {"3c": ""},
            {"3d": ""},
            {"3e": ""},
            {"3f": ""},
            {"3g": ""},
            {"3h": ""},
            {"3i": ""},
            {"3j": ""},
            {"3k": ""},
            {"3l": ""},
            {"3m": ""},
            {"plo_i": "sdfsf"},
            {"peo1_i": "0"},
            {"peo2_i": "0"},
            {"peo3_i": "0"},
            {"plo_e": "sdfs"},
            {"peo1_e": "0"},
            {"peo2_e": "0"},
            {"peo3_e": "0"},
            {"plo_d": "sdfs"},
            {"peo1_d": "0"},
            {"peo2_d": "0"},
            {"peo3_d": "0"},
            {"pea1_mission_a": "0"},
            {"pea1_mission_b": "0"},
            {"pea1_mission_c": "0"},
            {"pea1_mission_d": "0"},
            {"pea2_mission_a": "0"},
            {"pea2_mission_b": "0"},
            {"pea2_mission_c": "0"},
            {"pea2_mission_d": "0"},
            {"pea3_mission_a": "0"},
            {"pea3_mission_b": "0"},
            {"pea3_mission_c": "0"},
            {"pea3_mission_d": "0"},
            {"course-code-1": "dfgd"},
            {"plo-addressed-1": "dfgd"},
            {"reco-1": "dfgd"},
            {"course-code-2": ""},
            {"plo-addressed-2": ""},
            {"reco-2": ""},
            {"course-code-3": ""},
            {"plo-addressed-3": ""},
            {"reco-3": ""},
            {"reco-faculty-1": "dfgd"},
            {"actions-faculty-1": "dfg"},
            {"reco-faculty-2": ""},
            {"actions-faculty-2": ""},
            {"reco-faculty-3": ""},
            {"actions-faculty-3": ""},
            {"percentage-1": ""},
            {"prelim-exam-1": ""},
            {"exams1-1": "1"},
            {"projects1-1": "1"},
            {"presentations1-1": "1"},
            {"assignments1-1": "1"},
            {"midterm-exam-1": "dfg"},
            {"exams2-1": "1"},
            {"projects2-1": "1"},
            {"presentations2-1": "1"},
            {"assignments2-1": "1"},
            {"final-exam-1": "dfg"},
            {"exams3-1": "1"},
            {"projects3-1": "1"},
            {"presentations3-1": "1"},
            {"assignments3-1": "1"},
            {"compre-exam-1": "dg"},
            {"exams4-1": "1"},
            {"projects4-1": "1"},
            {"presentations4-1": "1"},
            {"assignments4-1": "1"},
            {"percentage-2": ""},
            {"prelim-exam-2": ""},
            {"exams1-2": "0"},
            {"projects1-2": "0"},
            {"presentations1-2": "0"},
            {"assignments1-2": "0"},
            {"midterm-exam-2": ""},
            {"exams2-2": "0"},
            {"projects2-2": "0"},
            {"presentations2-2": "0"},
            {"assignments2-2": "0"},
            {"final-exam-2": ""},
            {"exams3-2": "0"},
            {"projects3-2": "0"},
            {"presentations3-2": "0"},
            {"assignments3-2": "0"},
            {"compre-exam-2": ""},
            {"exams4-2": "0"},
            {"projects4-2": "0"},
            {"presentations4-2": "0"},
            {"assignments4-2": "0"},
            {"percentage-3": ""},
            {"prelim-exam-3": ""},
            {"exams1-3": "0"},
            {"projects1-3": "0"},
            {"presentations1-3": "0"},
            {"assignments1-3": "0"},
            {"midterm-exam-3": ""},
            {"exams2-3": "0"},
            {"projects2-3": "0"},
            {"presentations2-3": "0"},
            {"assignments2-3": "0"},
            {"final-exam-3": ""},
            {"exams3-3": "0"},
            {"projects3-3": "0"},
            {"presentations3-3": "0"},
            {"assignments3-3": "0"},
            {"compre-exam-3": ""},
            {"exams4-3": "0"},
            {"projects4-3": "0"},
            {"presentations4-3": "0"},
            {"assignments4-3": "0"},
            {"course-title-1": ""},
            {"course-title-2": ""},
            {"course-title-3": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "cam.docx")
        output_file_path = os.path.join(
            script_directory, f"cam-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        a = DownloadModel(user_id=user_id, record_id=lastid, form_name="matrix")
        a.save()

        end = time.time()
        print("End-Time: ", end)
        print("Time: ", end - start)

        response = FileResponse(open(output_file_path, "rb"))
        response["Content-Disposition"] = 'attachment; filename="matrix-output.docx"'
        return response
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def cam_update(request):
    if request.method == "POST":
        data = json.loads(request.body)

        class_record = data["class_record_auto_fetch_data"]
        lastid = data["updateid"]
        user_id = data["user_id"]

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"course_title": "a"},
            {"course_code": "a"},
            {"pre_requisite": "a"},
            {"prior_courses_with_similar_plo": "a"},
            {"academic_year": "a"},
            {"faculty": "a"},
            {"clo-1": "sdfsd"},
            {"1a": ""},
            {"1b": ""},
            {"1c": ""},
            {"1d": ""},
            {"1e": ""},
            {"1f": ""},
            {"1g": ""},
            {"1h": ""},
            {"1i": ""},
            {"1j": ""},
            {"1k": ""},
            {"1l": ""},
            {"1m": ""},
            {"clo-2": ""},
            {"2a": ""},
            {"2b": ""},
            {"2c": ""},
            {"2d": ""},
            {"2e": ""},
            {"2f": ""},
            {"2g": ""},
            {"2h": ""},
            {"2i": ""},
            {"2j": ""},
            {"2k": ""},
            {"2l": ""},
            {"2m": ""},
            {"clo-3": ""},
            {"3a": ""},
            {"3b": ""},
            {"3c": ""},
            {"3d": ""},
            {"3e": ""},
            {"3f": ""},
            {"3g": ""},
            {"3h": ""},
            {"3i": ""},
            {"3j": ""},
            {"3k": ""},
            {"3l": ""},
            {"3m": ""},
            {"plo_i": "sdfsf"},
            {"peo1_i": "0"},
            {"peo2_i": "0"},
            {"peo3_i": "0"},
            {"plo_e": "sdfs"},
            {"peo1_e": "0"},
            {"peo2_e": "0"},
            {"peo3_e": "0"},
            {"plo_d": "sdfs"},
            {"peo1_d": "0"},
            {"peo2_d": "0"},
            {"peo3_d": "0"},
            {"pea1_mission_a": "0"},
            {"pea1_mission_b": "0"},
            {"pea1_mission_c": "0"},
            {"pea1_mission_d": "0"},
            {"pea2_mission_a": "0"},
            {"pea2_mission_b": "0"},
            {"pea2_mission_c": "0"},
            {"pea2_mission_d": "0"},
            {"pea3_mission_a": "0"},
            {"pea3_mission_b": "0"},
            {"pea3_mission_c": "0"},
            {"pea3_mission_d": "0"},
            {"course-code-1": "dfgd"},
            {"plo-addressed-1": "dfgd"},
            {"reco-1": "dfgd"},
            {"course-code-2": ""},
            {"plo-addressed-2": ""},
            {"reco-2": ""},
            {"course-code-3": ""},
            {"plo-addressed-3": ""},
            {"reco-3": ""},
            {"reco-faculty-1": "dfgd"},
            {"actions-faculty-1": "dfg"},
            {"reco-faculty-2": ""},
            {"actions-faculty-2": ""},
            {"reco-faculty-3": ""},
            {"actions-faculty-3": ""},
            {"percentage-1": ""},
            {"prelim-exam-1": ""},
            {"exams1-1": "1"},
            {"projects1-1": "1"},
            {"presentations1-1": "1"},
            {"assignments1-1": "1"},
            {"midterm-exam-1": "dfg"},
            {"exams2-1": "1"},
            {"projects2-1": "1"},
            {"presentations2-1": "1"},
            {"assignments2-1": "1"},
            {"final-exam-1": "dfg"},
            {"exams3-1": "1"},
            {"projects3-1": "1"},
            {"presentations3-1": "1"},
            {"assignments3-1": "1"},
            {"compre-exam-1": "dg"},
            {"exams4-1": "1"},
            {"projects4-1": "1"},
            {"presentations4-1": "1"},
            {"assignments4-1": "1"},
            {"percentage-2": ""},
            {"prelim-exam-2": ""},
            {"exams1-2": "0"},
            {"projects1-2": "0"},
            {"presentations1-2": "0"},
            {"assignments1-2": "0"},
            {"midterm-exam-2": ""},
            {"exams2-2": "0"},
            {"projects2-2": "0"},
            {"presentations2-2": "0"},
            {"assignments2-2": "0"},
            {"final-exam-2": ""},
            {"exams3-2": "0"},
            {"projects3-2": "0"},
            {"presentations3-2": "0"},
            {"assignments3-2": "0"},
            {"compre-exam-2": ""},
            {"exams4-2": "0"},
            {"projects4-2": "0"},
            {"presentations4-2": "0"},
            {"assignments4-2": "0"},
            {"percentage-3": ""},
            {"prelim-exam-3": ""},
            {"exams1-3": "0"},
            {"projects1-3": "0"},
            {"presentations1-3": "0"},
            {"assignments1-3": "0"},
            {"midterm-exam-3": ""},
            {"exams2-3": "0"},
            {"projects2-3": "0"},
            {"presentations2-3": "0"},
            {"assignments2-3": "0"},
            {"final-exam-3": ""},
            {"exams3-3": "0"},
            {"projects3-3": "0"},
            {"presentations3-3": "0"},
            {"assignments3-3": "0"},
            {"compre-exam-3": ""},
            {"exams4-3": "0"},
            {"projects4-3": "0"},
            {"presentations4-3": "0"},
            {"assignments4-3": "0"},
            {"course-title-1": ""},
            {"course-title-2": ""},
            {"course-title-3": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "cam.docx")
        output_file_path = os.path.join(
            script_directory, f"cam-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        try:
            DownloadModel.objects.update_or_create(
                user_id=user_id, record_id=lastid, form_name="matrix", defaults={}
            )
        except DownloadModel.MultipleObjectsReturned:
            DownloadModel.objects.filter(
                user_id=user_id, record_id=lastid, form_name="matrix"
            ).delete()
            DownloadModel.objects.get_or_create(
                user_id=user_id, record_id=lastid, form_name="matrix"
            )

        return JsonResponse({"status": "success", "message": "File created"})

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def datasheet(request):
    if request.method == "POST":
        start = time.time()
        print("Start-Time: ", start)
        data = json.loads(request.body)

        class_record = data["class_record_auto_fetch_data"]
        user_id = data["user_id"]
        lastid = data["lastid"]
        obj, created = AutoFetchModel.objects.update_or_create(
            user_id=user_id, category="auto-fetch-data", defaults={"name": class_record}
        )

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"course_code": "a"},
            {"course_title": "a"},
            {"pre_requisite": "a"},
            {"prior_courses_with_similar_plo": "a"},
            {"academic_year": "a"},
            {"faculty": "a"},
            {"frequency_91100": ""},
            {"percentage_91100": ""},
            {"frequency_8190": ""},
            {"percentage_8190": ""},
            {"frequency_7180": ""},
            {"percentage_7180": ""},
            {"frequency_6170": ""},
            {"percentage_6170": ""},
            {"frequency_5160": ""},
            {"percentage_5160": ""},
            {"frequency_4150": ""},
            {"percentage_4150": ""},
            {"frequency_3140": ""},
            {"percentage_3140": ""},
            {"frequency_2130": ""},
            {"percentage_2130": ""},
            {"frequency_1120": ""},
            {"percentage_1120": ""},
            {"frequency_110": ""},
            {"percentage_110": ""},
            {"frequency_total": ""},
            {"percentage_total": ""},
            {"description_clo1": ""},
            {"description_clo2": ""},
            {"description_clo3": ""},
            {"items_1": ""},
            {"frequency_1": ""},
            {"percentage_1": ""},
            {"items_2": ""},
            {"frequency_2": ""},
            {"percentage_2": ""},
            {"items_3": ""},
            {"frequency_3": ""},
            {"percentage_3": ""},
            {"items_4": ""},
            {"frequency_4": ""},
            {"percentage_4": ""},
            {"items_5": ""},
            {"frequency_5": ""},
            {"percentage_5": ""},
            {"items_6": ""},
            {"frequency_6": ""},
            {"percentage_6": ""},
            {"items_7": ""},
            {"frequency_7": ""},
            {"percentage_7": ""},
            {"items_8": ""},
            {"frequency_8": ""},
            {"percentage_8": ""},
            {"items_9": ""},
            {"frequency_9": ""},
            {"percentage_9": ""},
            {"items_10": ""},
            {"frequency_10": ""},
            {"percentage_10": ""},
            {"items_11": ""},
            {"frequency_11": ""},
            {"percentage_11": ""},
            {"items_12": ""},
            {"frequency_12": ""},
            {"percentage_12": ""},
            {"items_13": ""},
            {"frequency_13": ""},
            {"percentage_13": ""},
            {"items_14": ""},
            {"frequency_14": ""},
            {"percentage_14": ""},
            {"items_15": ""},
            {"frequency_15": ""},
            {"percentage_15": ""},
            {"items_16": ""},
            {"frequency_16": ""},
            {"percentage_16": ""},
            {"items_17": ""},
            {"frequency_17": ""},
            {"percentage_17": ""},
            {"items_18": ""},
            {"frequency_18": ""},
            {"percentage_18": ""},
            {"items_19": ""},
            {"frequency_19": ""},
            {"percentage_19": ""},
            {"items_20": ""},
            {"frequency_20": ""},
            {"percentage_20": ""},
            {"items_21": ""},
            {"frequency_21": ""},
            {"percentage_21": ""},
            {"items_22": ""},
            {"frequency_22": ""},
            {"percentage_22": ""},
            {"items_23": ""},
            {"frequency_23": ""},
            {"percentage_23": ""},
            {"items_24": ""},
            {"frequency_24": ""},
            {"percentage_24": ""},
            {"items_25": ""},
            {"frequency_25": ""},
            {"percentage_25": ""},
            {"description_itemsclo1": ""},
            {"items_26": ""},
            {"frequency_26": ""},
            {"percentage_26": ""},
            {"items_27": ""},
            {"frequency_27": ""},
            {"percentage_27": ""},
            {"items_28": ""},
            {"frequency_28": ""},
            {"percentage_28": ""},
            {"items_29": ""},
            {"frequency_29": ""},
            {"percentage_29": ""},
            {"items_30": ""},
            {"frequency_30": ""},
            {"percentage_30": ""},
            {"items_31": ""},
            {"frequency_31": ""},
            {"percentage_31": ""},
            {"items_32": ""},
            {"frequency_32": ""},
            {"percentage_32": ""},
            {"items_33": ""},
            {"frequency_33": ""},
            {"percentage_33": ""},
            {"items_34": ""},
            {"frequency_34": ""},
            {"percentage_34": ""},
            {"items_35": ""},
            {"frequency_35": ""},
            {"percentage_35": ""},
            {"items_36": ""},
            {"frequency_36": ""},
            {"percentage_36": ""},
            {"items_37": ""},
            {"frequency_37": ""},
            {"percentage_37": ""},
            {"items_38": ""},
            {"frequency_38": ""},
            {"percentage_38": ""},
            {"items_39": ""},
            {"frequency_39": ""},
            {"percentage_39": ""},
            {"items_40": ""},
            {"frequency_40": ""},
            {"percentage_40": ""},
            {"items_41": ""},
            {"frequency_41": ""},
            {"percentage_41": ""},
            {"items_42": ""},
            {"frequency_42": ""},
            {"percentage_42": ""},
            {"items_43": ""},
            {"frequency_43": ""},
            {"percentage_43": ""},
            {"items_44": ""},
            {"frequency_44": ""},
            {"percentage_44": ""},
            {"items_45": ""},
            {"frequency_45": ""},
            {"percentage_45": ""},
            {"items_46": ""},
            {"frequency_46": ""},
            {"percentage_46": ""},
            {"items_47": ""},
            {"frequency_47": ""},
            {"percentage_47": ""},
            {"items_48": ""},
            {"frequency_48": ""},
            {"percentage_48": ""},
            {"items_49": ""},
            {"frequency_49": ""},
            {"percentage_49": ""},
            {"items_50": ""},
            {"frequency_50": ""},
            {"percentage_50": ""},
            {"description_itemsclo2": ""},
            {"items_51": ""},
            {"frequency_51": ""},
            {"percentage_51": ""},
            {"items_52": ""},
            {"frequency_52": ""},
            {"percentage_52": ""},
            {"items_53": ""},
            {"frequency_53": ""},
            {"percentage_53": ""},
            {"items_54": ""},
            {"frequency_54": ""},
            {"percentage_54": ""},
            {"items_55": ""},
            {"frequency_55": ""},
            {"percentage_55": ""},
            {"items_56": ""},
            {"frequency_56": ""},
            {"percentage_56": ""},
            {"items_57": ""},
            {"frequency_57": ""},
            {"percentage_57": ""},
            {"items_58": ""},
            {"frequency_58": ""},
            {"percentage_58": ""},
            {"items_59": ""},
            {"frequency_59": ""},
            {"percentage_59": ""},
            {"items_60": ""},
            {"frequency_60": ""},
            {"percentage_60": ""},
            {"items_61": ""},
            {"frequency_61": ""},
            {"percentage_61": ""},
            {"items_62": ""},
            {"frequency_62": ""},
            {"percentage_62": ""},
            {"items_63": ""},
            {"frequency_63": ""},
            {"percentage_63": ""},
            {"items_64": ""},
            {"frequency_64": ""},
            {"percentage_64": ""},
            {"items_65": ""},
            {"frequency_65": ""},
            {"percentage_65": ""},
            {"items_66": ""},
            {"frequency_66": ""},
            {"percentage_66": ""},
            {"items_67": ""},
            {"frequency_67": ""},
            {"percentage_67": ""},
            {"items_68": ""},
            {"frequency_68": ""},
            {"percentage_68": ""},
            {"items_69": ""},
            {"frequency_69": ""},
            {"percentage_69": ""},
            {"items_70": ""},
            {"frequency_70": ""},
            {"percentage_70": ""},
            {"items_71": ""},
            {"frequency_71": ""},
            {"percentage_71": ""},
            {"items_72": ""},
            {"frequency_72": ""},
            {"percentage_72": ""},
            {"items_73": ""},
            {"frequency_73": ""},
            {"percentage_73": ""},
            {"items_74": ""},
            {"frequency_74": ""},
            {"percentage_74": ""},
            {"items_75": ""},
            {"frequency_75": ""},
            {"percentage_75": ""},
            {"description_itemsclo3": ""},
            {"description_100": ""},
            {"frequency_100": ""},
            {"percentage_100": ""},
            {"description_125": ""},
            {"frequency_125": ""},
            {"percentage_125": ""},
            {"description_150": ""},
            {"frequency_150": ""},
            {"percentage_150": ""},
            {"description_175": ""},
            {"frequency_175": ""},
            {"percentage_175": ""},
            {"description_200": ""},
            {"frequency_200": ""},
            {"percentage_200": ""},
            {"description_225": ""},
            {"frequency_225": ""},
            {"percentage_225": ""},
            {"description_250": ""},
            {"frequency_250": ""},
            {"percentage_250": ""},
            {"description_275": ""},
            {"frequency_275": ""},
            {"percentage_275": ""},
            {"description_300": ""},
            {"frequency_300": ""},
            {"percentage_300": ""},
            {"description_500": ""},
            {"frequency_500": ""},
            {"percentage_500": ""},
            {"description_od": ""},
            {"frequency_od": ""},
            {"percentage_od": ""},
            {"description_ud": ""},
            {"frequency_ud": ""},
            {"percentage_ud": ""},
            {"description_semestral": ""},
            {"total_direct_clo1": ""},
            {"total_direct_clo2": ""},
            {"total_direct_clo3": ""},
            {"total_indirect_clo1": ""},
            {"total_indirect_clo2": ""},
            {"total_indirect_clo3": ""},
            {"total_attainment_clo1": ""},
            {"total_attainment_clo2": ""},
            {"total_attainment_clo3": ""},
            {"total_attainment_clo3": ""},
            {"weight_clo1": ""},
            {"direct_clo1": ""},
            {"indirect_clo1": ""},
            {"attainment_clo1": ""},
            {"remarks_clo1": ""},
            {"weight_clo2": ""},
            {"direct_clo2": ""},
            {"indirect_clo2": ""},
            {"attainment_clo2": ""},
            {"remarks_clo2": ""},
            {"weight_clo3": ""},
            {"direct_clo3": ""},
            {"indirect_clo3": ""},
            {"attainment_clo3": ""},
            {"remarks_clo3": ""},
            {"direct_total": ""},
            {"indirect_total": ""},
            {"attainment_total": ""},
            {"remarks_total": ""},
            {"description_coo": ""},
            {"assessment_item1": ""},
            {"average_attainment1": ""},
            {"average_plo1": ""},
            {"average_peo1": ""},
            {"criteria_peo1": ""},
            {"criteria_plo1": ""},
            {"criteria_clo1": ""},
            {"assessment_item2": ""},
            {"average_attainment2": ""},
            {"average_plo2": ""},
            {"average_peo2": ""},
            {"criteria_peo2": ""},
            {"criteria_plo2": ""},
            {"criteria_clo2": ""},
            {"assessment_item3": ""},
            {"average_attainment3": ""},
            {"average_plo3": ""},
            {"average_peo3": ""},
            {"criteria_peo3": ""},
            {"criteria_plo3": ""},
            {"criteria_clo3": ""},
            {"assessment_item4": ""},
            {"average_attainment4": ""},
            {"average_plo4": ""},
            {"average_peo4": ""},
            {"criteria_peo4": ""},
            {"criteria_plo4": ""},
            {"criteria_clo4": ""},
            {"assessment_item5": ""},
            {"average_attainment5": ""},
            {"average_plo5": ""},
            {"average_peo5": ""},
            {"criteria_peo5": ""},
            {"criteria_plo5": ""},
            {"criteria_clo5": ""},
            {"assessment_item6": ""},
            {"average_attainment6": ""},
            {"average_plo6": ""},
            {"average_peo6": ""},
            {"criteria_peo6": ""},
            {"criteria_plo6": ""},
            {"criteria_clo6": ""},
            {"assessment_item7": ""},
            {"average_attainment7": ""},
            {"average_plo7": ""},
            {"average_peo7": ""},
            {"description_ppc": ""},
            {"clo1_sa": ""},
            {"clo1_a": ""},
            {"clo1_n": ""},
            {"clo1_d": ""},
            {"clo1_sd": ""},
            {"clo1_total": ""},
            {"clo1_tsa": ""},
            {"clo1_ta": ""},
            {"clo1_tn": ""},
            {"clo1_td": ""},
            {"clo1_tsd": ""},
            {"clo1_ttotal": ""},
            {"clo1_wm": ""},
            {"clo2_sa": ""},
            {"clo2_a": ""},
            {"clo2_n": ""},
            {"clo2_d": ""},
            {"clo2_sd": ""},
            {"clo2_total": ""},
            {"clo2_tsa": ""},
            {"clo2_ta": ""},
            {"clo2_tn": ""},
            {"clo2_td": ""},
            {"clo2_tsd": ""},
            {"clo2_ttotal": ""},
            {"clo2_wm": ""},
            {"clo3_sa": ""},
            {"clo3_a": ""},
            {"clo3_n": ""},
            {"clo3_d": ""},
            {"clo3_sd": ""},
            {"clo3_total": ""},
            {"clo3_tsa": ""},
            {"clo3_ta": ""},
            {"clo3_tn": ""},
            {"clo3_td": ""},
            {"clo3_tsd": ""},
            {"clo3_ttotal": ""},
            {"clo3_wm": ""},
            {"rate_clo1_1": "0"},
            {"rate_clo1_2": "0"},
            {"rate_clo1_3": "0"},
            {"rate_clo1_4": "0"},
            {"rate_clo2_1": "0"},
            {"rate_clo2_2": "0"},
            {"rate_clo2_3": "0"},
            {"rate_clo2_4": "0"},
            {"rate_clo3_1": "0"},
            {"rate_clo3_2": "0"},
            {"rate_clo3_3": "0"},
            {"rate_clo3_4": "0"},
            {"description_indirectform": ""},
            {"prev_reco": ""},
            {"faculty_suggestion": ""},
            {"students_suggestion": ""},
            {"id": ""},
            {"form_id": ""},
        ]
        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "datasheet.docx")
        output_file_path = os.path.join(
            script_directory, f"datasheet-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        a = DownloadModel(user_id=user_id, record_id=lastid, form_name="datasheet")
        a.save()

        end = time.time()
        print("End-Time: ", end)
        print("Time: ", end - start)

        response = FileResponse(open(output_file_path, "rb"))
        response["Content-Disposition"] = 'attachment; filename="datasheet-output.docx"'
        return response

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def datasheet_update(request):
    if request.method == "POST":
        data = json.loads(request.body)

        class_record = data["class_record_auto_fetch_data"]
        user_id = data["user_id"]
        lastid = data["updateid"]

        DownloadModel.objects.filter(
            user_id=user_id, record_id=lastid, form_name="datasheet"
        ).delete()

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"course_code": "a"},
            {"course_title": "a"},
            {"pre_requisite": "a"},
            {"prior_courses_with_similar_plo": "a"},
            {"academic_year": "a"},
            {"faculty": "a"},
            {"frequency_91100": ""},
            {"percentage_91100": ""},
            {"frequency_8190": ""},
            {"percentage_8190": ""},
            {"frequency_7180": ""},
            {"percentage_7180": ""},
            {"frequency_6170": ""},
            {"percentage_6170": ""},
            {"frequency_5160": ""},
            {"percentage_5160": ""},
            {"frequency_4150": ""},
            {"percentage_4150": ""},
            {"frequency_3140": ""},
            {"percentage_3140": ""},
            {"frequency_2130": ""},
            {"percentage_2130": ""},
            {"frequency_1120": ""},
            {"percentage_1120": ""},
            {"frequency_110": ""},
            {"percentage_110": ""},
            {"frequency_total": ""},
            {"percentage_total": ""},
            {"description_clo1": ""},
            {"description_clo2": ""},
            {"description_clo3": ""},
            {"items_1": ""},
            {"frequency_1": ""},
            {"percentage_1": ""},
            {"items_2": ""},
            {"frequency_2": ""},
            {"percentage_2": ""},
            {"items_3": ""},
            {"frequency_3": ""},
            {"percentage_3": ""},
            {"items_4": ""},
            {"frequency_4": ""},
            {"percentage_4": ""},
            {"items_5": ""},
            {"frequency_5": ""},
            {"percentage_5": ""},
            {"items_6": ""},
            {"frequency_6": ""},
            {"percentage_6": ""},
            {"items_7": ""},
            {"frequency_7": ""},
            {"percentage_7": ""},
            {"items_8": ""},
            {"frequency_8": ""},
            {"percentage_8": ""},
            {"items_9": ""},
            {"frequency_9": ""},
            {"percentage_9": ""},
            {"items_10": ""},
            {"frequency_10": ""},
            {"percentage_10": ""},
            {"items_11": ""},
            {"frequency_11": ""},
            {"percentage_11": ""},
            {"items_12": ""},
            {"frequency_12": ""},
            {"percentage_12": ""},
            {"items_13": ""},
            {"frequency_13": ""},
            {"percentage_13": ""},
            {"items_14": ""},
            {"frequency_14": ""},
            {"percentage_14": ""},
            {"items_15": ""},
            {"frequency_15": ""},
            {"percentage_15": ""},
            {"items_16": ""},
            {"frequency_16": ""},
            {"percentage_16": ""},
            {"items_17": ""},
            {"frequency_17": ""},
            {"percentage_17": ""},
            {"items_18": ""},
            {"frequency_18": ""},
            {"percentage_18": ""},
            {"items_19": ""},
            {"frequency_19": ""},
            {"percentage_19": ""},
            {"items_20": ""},
            {"frequency_20": ""},
            {"percentage_20": ""},
            {"items_21": ""},
            {"frequency_21": ""},
            {"percentage_21": ""},
            {"items_22": ""},
            {"frequency_22": ""},
            {"percentage_22": ""},
            {"items_23": ""},
            {"frequency_23": ""},
            {"percentage_23": ""},
            {"items_24": ""},
            {"frequency_24": ""},
            {"percentage_24": ""},
            {"items_25": ""},
            {"frequency_25": ""},
            {"percentage_25": ""},
            {"description_itemsclo1": ""},
            {"items_26": ""},
            {"frequency_26": ""},
            {"percentage_26": ""},
            {"items_27": ""},
            {"frequency_27": ""},
            {"percentage_27": ""},
            {"items_28": ""},
            {"frequency_28": ""},
            {"percentage_28": ""},
            {"items_29": ""},
            {"frequency_29": ""},
            {"percentage_29": ""},
            {"items_30": ""},
            {"frequency_30": ""},
            {"percentage_30": ""},
            {"items_31": ""},
            {"frequency_31": ""},
            {"percentage_31": ""},
            {"items_32": ""},
            {"frequency_32": ""},
            {"percentage_32": ""},
            {"items_33": ""},
            {"frequency_33": ""},
            {"percentage_33": ""},
            {"items_34": ""},
            {"frequency_34": ""},
            {"percentage_34": ""},
            {"items_35": ""},
            {"frequency_35": ""},
            {"percentage_35": ""},
            {"items_36": ""},
            {"frequency_36": ""},
            {"percentage_36": ""},
            {"items_37": ""},
            {"frequency_37": ""},
            {"percentage_37": ""},
            {"items_38": ""},
            {"frequency_38": ""},
            {"percentage_38": ""},
            {"items_39": ""},
            {"frequency_39": ""},
            {"percentage_39": ""},
            {"items_40": ""},
            {"frequency_40": ""},
            {"percentage_40": ""},
            {"items_41": ""},
            {"frequency_41": ""},
            {"percentage_41": ""},
            {"items_42": ""},
            {"frequency_42": ""},
            {"percentage_42": ""},
            {"items_43": ""},
            {"frequency_43": ""},
            {"percentage_43": ""},
            {"items_44": ""},
            {"frequency_44": ""},
            {"percentage_44": ""},
            {"items_45": ""},
            {"frequency_45": ""},
            {"percentage_45": ""},
            {"items_46": ""},
            {"frequency_46": ""},
            {"percentage_46": ""},
            {"items_47": ""},
            {"frequency_47": ""},
            {"percentage_47": ""},
            {"items_48": ""},
            {"frequency_48": ""},
            {"percentage_48": ""},
            {"items_49": ""},
            {"frequency_49": ""},
            {"percentage_49": ""},
            {"items_50": ""},
            {"frequency_50": ""},
            {"percentage_50": ""},
            {"description_itemsclo2": ""},
            {"items_51": ""},
            {"frequency_51": ""},
            {"percentage_51": ""},
            {"items_52": ""},
            {"frequency_52": ""},
            {"percentage_52": ""},
            {"items_53": ""},
            {"frequency_53": ""},
            {"percentage_53": ""},
            {"items_54": ""},
            {"frequency_54": ""},
            {"percentage_54": ""},
            {"items_55": ""},
            {"frequency_55": ""},
            {"percentage_55": ""},
            {"items_56": ""},
            {"frequency_56": ""},
            {"percentage_56": ""},
            {"items_57": ""},
            {"frequency_57": ""},
            {"percentage_57": ""},
            {"items_58": ""},
            {"frequency_58": ""},
            {"percentage_58": ""},
            {"items_59": ""},
            {"frequency_59": ""},
            {"percentage_59": ""},
            {"items_60": ""},
            {"frequency_60": ""},
            {"percentage_60": ""},
            {"items_61": ""},
            {"frequency_61": ""},
            {"percentage_61": ""},
            {"items_62": ""},
            {"frequency_62": ""},
            {"percentage_62": ""},
            {"items_63": ""},
            {"frequency_63": ""},
            {"percentage_63": ""},
            {"items_64": ""},
            {"frequency_64": ""},
            {"percentage_64": ""},
            {"items_65": ""},
            {"frequency_65": ""},
            {"percentage_65": ""},
            {"items_66": ""},
            {"frequency_66": ""},
            {"percentage_66": ""},
            {"items_67": ""},
            {"frequency_67": ""},
            {"percentage_67": ""},
            {"items_68": ""},
            {"frequency_68": ""},
            {"percentage_68": ""},
            {"items_69": ""},
            {"frequency_69": ""},
            {"percentage_69": ""},
            {"items_70": ""},
            {"frequency_70": ""},
            {"percentage_70": ""},
            {"items_71": ""},
            {"frequency_71": ""},
            {"percentage_71": ""},
            {"items_72": ""},
            {"frequency_72": ""},
            {"percentage_72": ""},
            {"items_73": ""},
            {"frequency_73": ""},
            {"percentage_73": ""},
            {"items_74": ""},
            {"frequency_74": ""},
            {"percentage_74": ""},
            {"items_75": ""},
            {"frequency_75": ""},
            {"percentage_75": ""},
            {"description_itemsclo3": ""},
            {"description_100": ""},
            {"frequency_100": ""},
            {"percentage_100": ""},
            {"description_125": ""},
            {"frequency_125": ""},
            {"percentage_125": ""},
            {"description_150": ""},
            {"frequency_150": ""},
            {"percentage_150": ""},
            {"description_175": ""},
            {"frequency_175": ""},
            {"percentage_175": ""},
            {"description_200": ""},
            {"frequency_200": ""},
            {"percentage_200": ""},
            {"description_225": ""},
            {"frequency_225": ""},
            {"percentage_225": ""},
            {"description_250": ""},
            {"frequency_250": ""},
            {"percentage_250": ""},
            {"description_275": ""},
            {"frequency_275": ""},
            {"percentage_275": ""},
            {"description_300": ""},
            {"frequency_300": ""},
            {"percentage_300": ""},
            {"description_500": ""},
            {"frequency_500": ""},
            {"percentage_500": ""},
            {"description_od": ""},
            {"frequency_od": ""},
            {"percentage_od": ""},
            {"description_ud": ""},
            {"frequency_ud": ""},
            {"percentage_ud": ""},
            {"description_semestral": ""},
            {"total_direct_clo1": ""},
            {"total_direct_clo2": ""},
            {"total_direct_clo3": ""},
            {"total_indirect_clo1": ""},
            {"total_indirect_clo2": ""},
            {"total_indirect_clo3": ""},
            {"total_attainment_clo1": ""},
            {"total_attainment_clo2": ""},
            {"total_attainment_clo3": ""},
            {"total_attainment_clo3": ""},
            {"weight_clo1": ""},
            {"direct_clo1": ""},
            {"indirect_clo1": ""},
            {"attainment_clo1": ""},
            {"remarks_clo1": ""},
            {"weight_clo2": ""},
            {"direct_clo2": ""},
            {"indirect_clo2": ""},
            {"attainment_clo2": ""},
            {"remarks_clo2": ""},
            {"weight_clo3": ""},
            {"direct_clo3": ""},
            {"indirect_clo3": ""},
            {"attainment_clo3": ""},
            {"remarks_clo3": ""},
            {"direct_total": ""},
            {"indirect_total": ""},
            {"attainment_total": ""},
            {"remarks_total": ""},
            {"description_coo": ""},
            {"assessment_item1": ""},
            {"average_attainment1": ""},
            {"average_plo1": ""},
            {"average_peo1": ""},
            {"criteria_peo1": ""},
            {"criteria_plo1": ""},
            {"criteria_clo1": ""},
            {"assessment_item2": ""},
            {"average_attainment2": ""},
            {"average_plo2": ""},
            {"average_peo2": ""},
            {"criteria_peo2": ""},
            {"criteria_plo2": ""},
            {"criteria_clo2": ""},
            {"assessment_item3": ""},
            {"average_attainment3": ""},
            {"average_plo3": ""},
            {"average_peo3": ""},
            {"criteria_peo3": ""},
            {"criteria_plo3": ""},
            {"criteria_clo3": ""},
            {"assessment_item4": ""},
            {"average_attainment4": ""},
            {"average_plo4": ""},
            {"average_peo4": ""},
            {"criteria_peo4": ""},
            {"criteria_plo4": ""},
            {"criteria_clo4": ""},
            {"assessment_item5": ""},
            {"average_attainment5": ""},
            {"average_plo5": ""},
            {"average_peo5": ""},
            {"criteria_peo5": ""},
            {"criteria_plo5": ""},
            {"criteria_clo5": ""},
            {"assessment_item6": ""},
            {"average_attainment6": ""},
            {"average_plo6": ""},
            {"average_peo6": ""},
            {"criteria_peo6": ""},
            {"criteria_plo6": ""},
            {"criteria_clo6": ""},
            {"assessment_item7": ""},
            {"average_attainment7": ""},
            {"average_plo7": ""},
            {"average_peo7": ""},
            {"description_ppc": ""},
            {"clo1_sa": ""},
            {"clo1_a": ""},
            {"clo1_n": ""},
            {"clo1_d": ""},
            {"clo1_sd": ""},
            {"clo1_total": ""},
            {"clo1_tsa": ""},
            {"clo1_ta": ""},
            {"clo1_tn": ""},
            {"clo1_td": ""},
            {"clo1_tsd": ""},
            {"clo1_ttotal": ""},
            {"clo1_wm": ""},
            {"clo2_sa": ""},
            {"clo2_a": ""},
            {"clo2_n": ""},
            {"clo2_d": ""},
            {"clo2_sd": ""},
            {"clo2_total": ""},
            {"clo2_tsa": ""},
            {"clo2_ta": ""},
            {"clo2_tn": ""},
            {"clo2_td": ""},
            {"clo2_tsd": ""},
            {"clo2_ttotal": ""},
            {"clo2_wm": ""},
            {"clo3_sa": ""},
            {"clo3_a": ""},
            {"clo3_n": ""},
            {"clo3_d": ""},
            {"clo3_sd": ""},
            {"clo3_total": ""},
            {"clo3_tsa": ""},
            {"clo3_ta": ""},
            {"clo3_tn": ""},
            {"clo3_td": ""},
            {"clo3_tsd": ""},
            {"clo3_ttotal": ""},
            {"clo3_wm": ""},
            {"rate_clo1_1": "0"},
            {"rate_clo1_2": "0"},
            {"rate_clo1_3": "0"},
            {"rate_clo1_4": "0"},
            {"rate_clo2_1": "0"},
            {"rate_clo2_2": "0"},
            {"rate_clo2_3": "0"},
            {"rate_clo2_4": "0"},
            {"rate_clo3_1": "0"},
            {"rate_clo3_2": "0"},
            {"rate_clo3_3": "0"},
            {"rate_clo3_4": "0"},
            {"description_indirectform": ""},
            {"prev_reco": ""},
            {"faculty_suggestion": ""},
            {"students_suggestion": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "datasheet.docx")
        output_file_path = os.path.join(
            script_directory, f"datasheet-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        try:
            DownloadModel.objects.update_or_create(
                user_id=user_id, record_id=lastid, form_name="datasheet", defaults={}
            )
        except DownloadModel.MultipleObjectsReturned:
            DownloadModel.objects.filter(
                user_id=user_id, record_id=lastid, form_name="datasheet"
            ).delete()
            DownloadModel.objects.get_or_create(
                user_id=user_id, record_id=lastid, form_name="datasheet"
            )

        return JsonResponse({"status": "success", "message": "File created"})

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def course_assessment(request):
    if request.method == "POST":
        # Process incoming data
        start = time.time()
        data = json.loads(request.body)

        class_record = data["class_record_auto_fetch_data"]
        user_id = data["user_id"]
        lastid = data["lastid"]
        obj, created = AutoFetchModel.objects.update_or_create(
            user_id=user_id, category="auto-fetch-data", defaults={"name": class_record}
        )

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"semester": "a"},
            {"academic_year": "a"},
            {"sections-1": "dfgd"},
            {"course-code-1": "dfgd"},
            {"course-title-1": ""},
            {"plo-1": ""},
            {"assessment-tool-1": ""},
            {"tools-used-1": ""},
            {"students-1": "dfgd"},
            {"tool-passed-1": ""},
            {"course-passed-1": ""},
            {"exhibit-1": "dfgd"},
            {"sections-2": ""},
            {"course-code-2": ""},
            {"course-title-2": ""},
            {"plo-2": ""},
            {"assessment-tool-2": ""},
            {"tools-used-2": ""},
            {"students-2": ""},
            {"tool-passed-2": ""},
            {"course-passed-2": ""},
            {"exhibit-2": ""},
            {"sections-3": ""},
            {"course-code-3": ""},
            {"course-title-3": ""},
            {"plo-3": ""},
            {"assessment-tool-3": ""},
            {"tools-used-3": ""},
            {"students-3": ""},
            {"tool-passed-3": ""},
            {"course-passed-3": ""},
            {"exhibit-3": ""},
            {"experienced-1": ""},
            {"interventions-1": "dfgd"},
            {"practices-1": "dfgd"},
            {"improvement-1": "dfgd"},
            {"experienced-2": ""},
            {"interventions-2": ""},
            {"practices-2": ""},
            {"improvement-2": ""},
            {"experienced-3": ""},
            {"interventions-3": ""},
            {"practices-3": ""},
            {"improvement-3": ""},
            {"sections-4": ""},
            {"sections-5": ""},
            {"sections-6": ""},
            {"course-code-4": ""},
            {"course-code-5": ""},
            {"course-code-6": ""},
            {"course-title-4": ""},
            {"course-title-5": ""},
            {"course-title-6": ""},
            {"plo-4": ""},
            {"plo-5": ""},
            {"plo-6": ""},
            {"assessment-tool-4": ""},
            {"assessment-tool-5": ""},
            {"assessment-tool-6": ""},
            {"tools-used-4": ""},
            {"tools-used-5": ""},
            {"tools-used-6": ""},
            {"students-4": ""},
            {"students-5": ""},
            {"students-6": ""},
            {"tool-passed-4": ""},
            {"tool-passed-5": ""},
            {"tool-passed-6": ""},
            {"course-passed-4": ""},
            {"course-passed-5": ""},
            {"course-passed-6": ""},
            {"exhibit-4": ""},
            {"exhibit-5": ""},
            {"exhibit-6": ""},
            {"experienced-4": ""},
            {"experienced-5": ""},
            {"experienced-6": ""},
            {"interventions-4": ""},
            {"interventions-5": ""},
            {"interventions-6": ""},
            {"practices-4": ""},
            {"practices-5": ""},
            {"practices-6": ""},
            {"improvement-4": ""},
            {"improvement-5": ""},
            {"improvement-6": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        # inputvalues = [
        #     {"edit_values":"test values"},{"button_count":"test values"},{"auto_fetch_data":"test values"},{"sections-1":""},{"students-1":""},{"exhibit-1":""},{"sections-2":""},{"students-2":""},{"exhibit-2":""},{"sections-3":""},{"students-3":""},{"exhibit-3":""},{"interventions-1":""},{"practices-1":""},{"improvement-1":""},{"interventions-2":""},{"practices-2":""},{"improvement-2":""},{"interventions-3":""},{"practices-3":""},{"improvement-3":""},{"sections-4":""},{"sections-5":""},{"sections-6":""},{"students-4":""},{"students-5":""},{"students-6":""},{"exhibit-4":""},{"exhibit-5":""},{"exhibit-6":""},{"interventions-4":""},{"interventions-5":""},{"interventions-6":""},{"practices-4":""},{"practices-5":""},{"practices-6":""},{"improvement-4":""},{"improvement-5":""},{"improvement-6":""},{"id":""},{"form_id":""}
        # ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "cas.docx")
        output_file_path = os.path.join(
            script_directory, f"cas-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        a = DownloadModel(user_id=user_id, record_id=lastid, form_name="summary")
        a.save()

        end = time.time()
        print("End-Time: ", end)
        print("Time: ", end - start)

        response = FileResponse(open(output_file_path, "rb"))
        response["Content-Disposition"] = 'attachment; filename="course-output.docx"'
        return response

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def course_assessment_update(request):
    if request.method == "POST":
        # Process incoming data
        data = json.loads(request.body)

        user_id = data["user_id"]
        lastid = data["updateid"]

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"semester": "a"},
            {"academic_year": "a"},
            {"sections-1": "dfgd"},
            {"course-code-1": "dfgd"},
            {"course-title-1": ""},
            {"plo-1": ""},
            {"assessment-tool-1": ""},
            {"tools-used-1": ""},
            {"students-1": "dfgd"},
            {"tool-passed-1": ""},
            {"course-passed-1": ""},
            {"exhibit-1": "dfgd"},
            {"sections-2": ""},
            {"course-code-2": ""},
            {"course-title-2": ""},
            {"plo-2": ""},
            {"assessment-tool-2": ""},
            {"tools-used-2": ""},
            {"students-2": ""},
            {"tool-passed-2": ""},
            {"course-passed-2": ""},
            {"exhibit-2": ""},
            {"sections-3": ""},
            {"course-code-3": ""},
            {"course-title-3": ""},
            {"plo-3": ""},
            {"assessment-tool-3": ""},
            {"tools-used-3": ""},
            {"students-3": ""},
            {"tool-passed-3": ""},
            {"course-passed-3": ""},
            {"exhibit-3": ""},
            {"experienced-1": ""},
            {"interventions-1": "dfgd"},
            {"practices-1": "dfgd"},
            {"improvement-1": "dfgd"},
            {"experienced-2": ""},
            {"interventions-2": ""},
            {"practices-2": ""},
            {"improvement-2": ""},
            {"experienced-3": ""},
            {"interventions-3": ""},
            {"practices-3": ""},
            {"improvement-3": ""},
            {"sections-4": ""},
            {"sections-5": ""},
            {"sections-6": ""},
            {"course-code-4": ""},
            {"course-code-5": ""},
            {"course-code-6": ""},
            {"course-title-4": ""},
            {"course-title-5": ""},
            {"course-title-6": ""},
            {"plo-4": ""},
            {"plo-5": ""},
            {"plo-6": ""},
            {"assessment-tool-4": ""},
            {"assessment-tool-5": ""},
            {"assessment-tool-6": ""},
            {"tools-used-4": ""},
            {"tools-used-5": ""},
            {"tools-used-6": ""},
            {"students-4": ""},
            {"students-5": ""},
            {"students-6": ""},
            {"tool-passed-4": ""},
            {"tool-passed-5": ""},
            {"tool-passed-6": ""},
            {"course-passed-4": ""},
            {"course-passed-5": ""},
            {"course-passed-6": ""},
            {"exhibit-4": ""},
            {"exhibit-5": ""},
            {"exhibit-6": ""},
            {"experienced-4": ""},
            {"experienced-5": ""},
            {"experienced-6": ""},
            {"interventions-4": ""},
            {"interventions-5": ""},
            {"interventions-6": ""},
            {"practices-4": ""},
            {"practices-5": ""},
            {"practices-6": ""},
            {"improvement-4": ""},
            {"improvement-5": ""},
            {"improvement-6": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        # inputvalues = [
        #     {"edit_values":"test values"},{"button_count":"test values"},{"auto_fetch_data":"test values"},{"sections-1":""},{"students-1":""},{"exhibit-1":""},{"sections-2":""},{"students-2":""},{"exhibit-2":""},{"sections-3":""},{"students-3":""},{"exhibit-3":""},{"interventions-1":""},{"practices-1":""},{"improvement-1":""},{"interventions-2":""},{"practices-2":""},{"improvement-2":""},{"interventions-3":""},{"practices-3":""},{"improvement-3":""},{"sections-4":""},{"sections-5":""},{"sections-6":""},{"students-4":""},{"students-5":""},{"students-6":""},{"exhibit-4":""},{"exhibit-5":""},{"exhibit-6":""},{"interventions-4":""},{"interventions-5":""},{"interventions-6":""},{"practices-4":""},{"practices-5":""},{"practices-6":""},{"improvement-4":""},{"improvement-5":""},{"improvement-6":""},{"id":""},{"form_id":""}
        # ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "cas.docx")
        output_file_path = os.path.join(
            script_directory, f"cas-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        try:
            DownloadModel.objects.update_or_create(
                user_id=user_id, record_id=lastid, form_name="summary", defaults={}
            )
        except DownloadModel.MultipleObjectsReturned:
            DownloadModel.objects.filter(
                user_id=user_id, record_id=lastid, form_name="summary"
            ).delete()
            DownloadModel.objects.get_or_create(
                user_id=user_id, record_id=lastid, form_name="summary"
            )

        return JsonResponse({"status": "success", "message": "File created"})

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def tos(request):
    if request.method == "POST":
        # Process incoming data
        start = time.time()
        data = json.loads(request.body)
        class_record = data["class_record_auto_fetch_data"]
        user_id = data["user_id"]
        lastid = data["lastid"]
        obj, created = AutoFetchModel.objects.update_or_create(
            user_id=user_id, category="auto-fetch-data", defaults={"name": class_record}
        )

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"course_code": "s"},
            {"course_title": "s"},
            {"number_units": "s"},
            {"sections": "s"},
            {"exam": ""},
            {"semesters": "s"},
            {"faculty": ""},
            {"name-topic-1": ""},
            {"hours-instruction-1": ""},
            {"%-hours-1": ""},
            {"cilo-1": ""},
            {"1R": ""},
            {"1U": ""},
            {"1A1": ""},
            {"1A2": ""},
            {"1E": ""},
            {"1C": ""},
            {"total-1": ""},
            {"items-1": ""},
            {"%-items-1": ""},
            {"name-topic-2": ""},
            {"hours-instruction-2": ""},
            {"%-hours-2": ""},
            {"cilo-2": ""},
            {"2R": ""},
            {"2U": ""},
            {"2A1": ""},
            {"2A2": ""},
            {"2E": ""},
            {"2C": ""},
            {"total-2": ""},
            {"items-2": ""},
            {"%-items-2": ""},
            {"name-topic-3": ""},
            {"hours-instruction-3": ""},
            {"%-hours-3": ""},
            {"cilo-3": ""},
            {"3R": ""},
            {"3U": ""},
            {"3A1": ""},
            {"3A2": ""},
            {"3E": ""},
            {"3C": ""},
            {"total-3": ""},
            {"items-3": ""},
            {"%-items-3": ""},
            {"clo-1": ""},
            {"questionsexam-1": ""},
            {"level-taxonomy-1": ""},
            {"comments-1": ""},
            {"clo-2": ""},
            {"questionsexam-2": ""},
            {"level-taxonomy-2": ""},
            {"questions-clo-1": ""},
            {"questions-clo-2": ""},
            {"comments-2": ""},
            {"clo-3": ""},
            {"questionsexam-3": ""},
            {"level-taxonomy-3": ""},
            {"questions-clo-3": ""},
            {"comments-3": ""},
            {"prepared_by": ""},
            {"checked_by": ""},
            {"noted_by": ""},
            {"flexRadioDefault1y-1": ""},
            {"flexRadioDefault1y-2": ""},
            {"flexRadioDefault1y-3": ""},
            {"flexRadioDefault1n-1": ""},
            {"flexRadioDefault1n-2": ""},
            {"flexRadioDefault1n-3": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "tos.docx")
        output_file_path = os.path.join(
            script_directory, f"tos-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        a = DownloadModel(user_id=user_id, record_id=lastid, form_name="tos")
        a.save()

        end = time.time()
        print("End-Time: ", end)
        print("Time: ", end - start)

        response = FileResponse(open(output_file_path, "rb"))
        response["Content-Disposition"] = 'attachment; filename="tos-output.docx"'
        return response

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def tos_update(request):
    if request.method == "POST":
        # Process incoming data
        data = json.loads(request.body)
        user_id = data["user_id"]
        lastid = data["updateid"]

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"course_code": "s"},
            {"course_title": "s"},
            {"number_units": "s"},
            {"sections": "s"},
            {"exam": ""},
            {"semesters": "s"},
            {"faculty": ""},
            {"name-topic-1": ""},
            {"hours-instruction-1": ""},
            {"%-hours-1": ""},
            {"cilo-1": ""},
            {"1R": ""},
            {"1U": ""},
            {"1A1": ""},
            {"1A2": ""},
            {"1E": ""},
            {"1C": ""},
            {"total-1": ""},
            {"items-1": ""},
            {"%-items-1": ""},
            {"name-topic-2": ""},
            {"hours-instruction-2": ""},
            {"%-hours-2": ""},
            {"cilo-2": ""},
            {"2R": ""},
            {"2U": ""},
            {"2A1": ""},
            {"2A2": ""},
            {"2E": ""},
            {"2C": ""},
            {"total-2": ""},
            {"items-2": ""},
            {"%-items-2": ""},
            {"name-topic-3": ""},
            {"hours-instruction-3": ""},
            {"%-hours-3": ""},
            {"cilo-3": ""},
            {"3R": ""},
            {"3U": ""},
            {"3A1": ""},
            {"3A2": ""},
            {"3E": ""},
            {"3C": ""},
            {"total-3": ""},
            {"items-3": ""},
            {"%-items-3": ""},
            {"clo-1": ""},
            {"questionsexam-1": ""},
            {"level-taxonomy-1": ""},
            {"comments-1": ""},
            {"clo-2": ""},
            {"questionsexam-2": ""},
            {"level-taxonomy-2": ""},
            {"questions-clo-1": ""},
            {"questions-clo-2": ""},
            {"comments-2": ""},
            {"clo-3": ""},
            {"questionsexam-3": ""},
            {"level-taxonomy-3": ""},
            {"questions-clo-3": ""},
            {"comments-3": ""},
            {"prepared_by": ""},
            {"checked_by": ""},
            {"noted_by": ""},
            {"flexRadioDefault1y-1": ""},
            {"flexRadioDefault1y-2": ""},
            {"flexRadioDefault1y-3": ""},
            {"flexRadioDefault1n-1": ""},
            {"flexRadioDefault1n-2": ""},
            {"flexRadioDefault1n-3": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "tos.docx")
        output_file_path = os.path.join(
            script_directory, f"tos-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        try:
            DownloadModel.objects.update_or_create(
                user_id=user_id, record_id=lastid, form_name="tos", defaults={}
            )
        except DownloadModel.MultipleObjectsReturned:
            DownloadModel.objects.filter(
                user_id=user_id, record_id=lastid, form_name="tos"
            ).delete()
            DownloadModel.objects.get_or_create(
                user_id=user_id, record_id=lastid, form_name="tos"
            )

        return JsonResponse({"status": "success", "message": "File created"})

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def plo(request):
    if request.method == "POST":
        start = time.time()
        # Process incoming data
        data = json.loads(request.body)
        class_record = data["class_record_auto_fetch_data"]
        user_id = data["user_id"]
        lastid = data["lastid"]
        obj, created = AutoFetchModel.objects.update_or_create(
            user_id=user_id, category="auto-fetch-data", defaults={"name": class_record}
        )

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"prog_obj": ""},
            {"prog_desc": ""},
            {"ilo": ""},
            {"peo": ""},
            {"program": ""},
            {"college": ""},
            {"course-title-1": ""},
            {"course-title-2": ""},
            {"course-title-3": ""},
            {"course-title-4": ""},
            {"course-title-5": ""},
            {"course-title-6": ""},
            {"course-title-7": ""},
            {"course-title-8": ""},
            {"course-title-9": ""},
            {"course-title-10": ""},
            {"course-title-11": ""},
            {"course-title-12": ""},
            {"course-title-13": ""},
            {"course-title-14": ""},
            {"course-title-15": ""},
            {"plo-1": ""},
            {"plo-2": ""},
            {"plo-3": ""},
            {"plo-4": ""},
            {"plo-5": ""},
            {"plo-6": ""},
            {"plo-7": ""},
            {"plo-8": ""},
            {"plo-9": ""},
            {"plo-10": ""},
            {"plo-11": ""},
            {"plo-12": ""},
            {"plo-13": ""},
            {"plo-14": ""},
            {"plo-15": ""},
            {"assessment-tool-1": ""},
            {"assessment-tool-2": ""},
            {"assessment-tool-3": ""},
            {"assessment-tool-4": ""},
            {"assessment-tool-5": ""},
            {"assessment-tool-6": ""},
            {"assessment-tool-7": ""},
            {"assessment-tool-8": ""},
            {"assessment-tool-9": ""},
            {"assessment-tool-10": ""},
            {"assessment-tool-11": ""},
            {"assessment-tool-12": ""},
            {"assessment-tool-13": ""},
            {"assessment-tool-14": ""},
            {"assessment-tool-15": ""},
            {"tools-used-1": ""},
            {"tools-used-2": ""},
            {"tools-used-3": ""},
            {"tools-used-4": ""},
            {"tools-used-5": ""},
            {"tools-used-6": ""},
            {"tools-used-7": ""},
            {"tools-used-8": ""},
            {"tools-used-9": ""},
            {"tools-used-10": ""},
            {"tools-used-11": ""},
            {"tools-used-12": ""},
            {"tools-used-13": ""},
            {"tools-used-14": ""},
            {"tools-used-15": ""},
            {"students-1": ""},
            {"students-2": ""},
            {"students-3": ""},
            {"students-4": ""},
            {"students-5": ""},
            {"students-6": ""},
            {"students-7": ""},
            {"students-8": ""},
            {"students-9": ""},
            {"students-10": ""},
            {"students-11": ""},
            {"students-12": ""},
            {"students-13": ""},
            {"students-14": ""},
            {"students-15": ""},
            {"tool-passed-1": ""},
            {"tool-passed-2": ""},
            {"tool-passed-3": ""},
            {"tool-passed-4": ""},
            {"tool-passed-5": ""},
            {"tool-passed-6": ""},
            {"tool-passed-7": ""},
            {"tool-passed-8": ""},
            {"tool-passed-9": ""},
            {"tool-passed-10": ""},
            {"tool-passed-11": ""},
            {"tool-passed-12": ""},
            {"tool-passed-13": ""},
            {"tool-passed-14": ""},
            {"tool-passed-15": ""},
            {"course-passed-1": ""},
            {"course-passed-2": ""},
            {"course-passed-3": ""},
            {"course-passed-4": ""},
            {"course-passed-5": ""},
            {"course-passed-6": ""},
            {"course-passed-7": ""},
            {"course-passed-8": ""},
            {"course-passed-9": ""},
            {"course-passed-10": ""},
            {"course-passed-11": ""},
            {"course-passed-12": ""},
            {"course-passed-13": ""},
            {"course-passed-14": ""},
            {"course-passed-15": ""},
            {"input_1_0_pfm": ""},
            {"input_1_1_pfm": ""},
            {"input_1_2_pfm": ""},
            {"input_1_3_pfm": ""},
            {"input_1_4_pfm": ""},
            {"input_1_5_pfm": ""},
            {"input_1_6_pfm": ""},
            {"input_1_7_pfm": ""},
            {"input_1_8_pfm": ""},
            {"input_1_9_pfm": ""},
            {"input_1_10_pfm": ""},
            {"input_1_11_pfm": ""},
            {"input_1_12_pfm": ""},
            {"input_1_13_pfm": ""},
            {"input_1_14_pfm": ""},
            {"input_1_15_pfm": ""},
            {"input_2_0_pfm": ""},
            {"input_2_1_pfm": ""},
            {"input_2_2_pfm": ""},
            {"input_2_3_pfm": ""},
            {"input_2_4_pfm": ""},
            {"input_2_5_pfm": ""},
            {"input_2_6_pfm": ""},
            {"input_2_7_pfm": ""},
            {"input_2_8_pfm": ""},
            {"input_2_9_pfm": ""},
            {"input_2_10_pfm": ""},
            {"input_2_11_pfm": ""},
            {"input_2_12_pfm": ""},
            {"input_2_13_pfm": ""},
            {"input_2_14_pfm": ""},
            {"input_2_15_pfm": ""},
            {"input_3_0_pfm": ""},
            {"input_3_1_pfm": ""},
            {"input_3_2_pfm": ""},
            {"input_3_3_pfm": ""},
            {"input_3_4_pfm": ""},
            {"input_3_5_pfm": ""},
            {"input_3_6_pfm": ""},
            {"input_3_7_pfm": ""},
            {"input_3_8_pfm": ""},
            {"input_3_9_pfm": ""},
            {"input_3_10_pfm": ""},
            {"input_3_11_pfm": ""},
            {"input_3_12_pfm": ""},
            {"input_3_13_pfm": ""},
            {"input_3_14_pfm": ""},
            {"input_3_15_pfm": ""},
            {"input_4_0_pfm": ""},
            {"input_4_1_pfm": ""},
            {"input_4_2_pfm": ""},
            {"input_4_3_pfm": ""},
            {"input_4_4_pfm": ""},
            {"input_4_5_pfm": ""},
            {"input_4_6_pfm": ""},
            {"input_4_7_pfm": ""},
            {"input_4_8_pfm": ""},
            {"input_4_9_pfm": ""},
            {"input_4_10_pfm": ""},
            {"input_4_11_pfm": ""},
            {"input_4_12_pfm": ""},
            {"input_4_13_pfm": ""},
            {"input_4_14_pfm": ""},
            {"input_4_15_pfm": ""},
            {"input_5_0_pfm": ""},
            {"input_5_1_pfm": ""},
            {"input_5_2_pfm": ""},
            {"input_5_3_pfm": ""},
            {"input_5_4_pfm": ""},
            {"input_5_5_pfm": ""},
            {"input_5_6_pfm": ""},
            {"input_5_7_pfm": ""},
            {"input_5_8_pfm": ""},
            {"input_5_9_pfm": ""},
            {"input_5_10_pfm": ""},
            {"input_5_11_pfm": ""},
            {"input_5_12_pfm": ""},
            {"input_5_13_pfm": ""},
            {"input_5_14_pfm": ""},
            {"input_5_15_pfm": ""},
            {"input_6_0_pfm": ""},
            {"input_6_1_pfm": ""},
            {"input_6_2_pfm": ""},
            {"input_6_3_pfm": ""},
            {"input_6_4_pfm": ""},
            {"input_6_5_pfm": ""},
            {"input_6_6_pfm": ""},
            {"input_6_7_pfm": ""},
            {"input_6_8_pfm": ""},
            {"input_6_9_pfm": ""},
            {"input_6_10_pfm": ""},
            {"input_6_11_pfm": ""},
            {"input_6_12_pfm": ""},
            {"input_6_13_pfm": ""},
            {"input_6_14_pfm": ""},
            {"input_6_15_pfm": ""},
            {"input_7_0_pfm": ""},
            {"input_7_1_pfm": ""},
            {"input_7_2_pfm": ""},
            {"input_7_3_pfm": ""},
            {"input_7_4_pfm": ""},
            {"input_7_5_pfm": ""},
            {"input_7_6_pfm": ""},
            {"input_7_7_pfm": ""},
            {"input_7_8_pfm": ""},
            {"input_7_9_pfm": ""},
            {"input_7_10_pfm": ""},
            {"input_7_11_pfm": ""},
            {"input_7_12_pfm": ""},
            {"input_7_13_pfm": ""},
            {"input_7_14_pfm": ""},
            {"input_7_15_pfm": ""},
            {"input_8_0_pfm": ""},
            {"input_8_1_pfm": ""},
            {"input_8_2_pfm": ""},
            {"input_8_3_pfm": ""},
            {"input_8_4_pfm": ""},
            {"input_8_5_pfm": ""},
            {"input_8_6_pfm": ""},
            {"input_8_7_pfm": ""},
            {"input_8_8_pfm": ""},
            {"input_8_9_pfm": ""},
            {"input_8_10_pfm": ""},
            {"input_8_11_pfm": ""},
            {"input_8_12_pfm": ""},
            {"input_8_13_pfm": ""},
            {"input_8_14_pfm": ""},
            {"input_8_15_pfm": ""},
            {"input_9_0_pfm": ""},
            {"input_9_1_pfm": ""},
            {"input_9_2_pfm": ""},
            {"input_9_3_pfm": ""},
            {"input_9_4_pfm": ""},
            {"input_9_5_pfm": ""},
            {"input_9_6_pfm": ""},
            {"input_9_7_pfm": ""},
            {"input_9_8_pfm": ""},
            {"input_9_9_pfm": ""},
            {"input_9_10_pfm": ""},
            {"input_9_11_pfm": ""},
            {"input_9_12_pfm": ""},
            {"input_9_13_pfm": ""},
            {"input_9_14_pfm": ""},
            {"input_9_15_pfm": ""},
            {"input_10_0_pfm": ""},
            {"input_10_1_pfm": ""},
            {"input_10_2_pfm": ""},
            {"input_10_3_pfm": ""},
            {"input_10_4_pfm": ""},
            {"input_10_5_pfm": ""},
            {"input_10_6_pfm": ""},
            {"input_10_7_pfm": ""},
            {"input_10_8_pfm": ""},
            {"input_10_9_pfm": ""},
            {"input_10_10_pfm": ""},
            {"input_10_11_pfm": ""},
            {"input_10_12_pfm": ""},
            {"input_10_13_pfm": ""},
            {"input_10_14_pfm": ""},
            {"input_10_15_pfm": ""},
            {"input_11_0_pfm": ""},
            {"input_11_1_pfm": ""},
            {"input_11_2_pfm": ""},
            {"input_11_3_pfm": ""},
            {"input_11_4_pfm": ""},
            {"input_11_5_pfm": ""},
            {"input_11_6_pfm": ""},
            {"input_11_7_pfm": ""},
            {"input_11_8_pfm": ""},
            {"input_11_9_pfm": ""},
            {"input_11_10_pfm": ""},
            {"input_11_11_pfm": ""},
            {"input_11_12_pfm": ""},
            {"input_11_13_pfm": ""},
            {"input_11_14_pfm": ""},
            {"input_11_15_pfm": ""},
            {"input_12_0_pfm": ""},
            {"input_12_1_pfm": ""},
            {"input_12_2_pfm": ""},
            {"input_12_3_pfm": ""},
            {"input_12_4_pfm": ""},
            {"input_12_5_pfm": ""},
            {"input_12_6_pfm": ""},
            {"input_12_7_pfm": ""},
            {"input_12_8_pfm": ""},
            {"input_12_9_pfm": ""},
            {"input_12_10_pfm": ""},
            {"input_12_11_pfm": ""},
            {"input_12_12_pfm": ""},
            {"input_12_13_pfm": ""},
            {"input_12_14_pfm": ""},
            {"input_12_15_pfm": ""},
            {"input_13_0_pfm": ""},
            {"input_13_1_pfm": ""},
            {"input_13_2_pfm": ""},
            {"input_13_3_pfm": ""},
            {"input_13_4_pfm": ""},
            {"input_13_5_pfm": ""},
            {"input_13_6_pfm": ""},
            {"input_13_7_pfm": ""},
            {"input_13_8_pfm": ""},
            {"input_13_9_pfm": ""},
            {"input_13_10_pfm": ""},
            {"input_13_11_pfm": ""},
            {"input_13_12_pfm": ""},
            {"input_13_13_pfm": ""},
            {"input_13_14_pfm": ""},
            {"input_13_15_pfm": ""},
            {"input_14_0_pfm": ""},
            {"input_14_1_pfm": ""},
            {"input_14_2_pfm": ""},
            {"input_14_3_pfm": ""},
            {"input_14_4_pfm": ""},
            {"input_14_5_pfm": ""},
            {"input_14_6_pfm": ""},
            {"input_14_7_pfm": ""},
            {"input_14_8_pfm": ""},
            {"input_14_9_pfm": ""},
            {"input_14_10_pfm": ""},
            {"input_14_11_pfm": ""},
            {"input_14_12_pfm": ""},
            {"input_14_13_pfm": ""},
            {"input_14_14_pfm": ""},
            {"input_14_15_pfm": ""},
            {"input_15_0_pfm": ""},
            {"input_15_1_pfm": ""},
            {"input_15_2_pfm": ""},
            {"input_15_3_pfm": ""},
            {"input_15_4_pfm": ""},
            {"input_15_5_pfm": ""},
            {"input_15_6_pfm": ""},
            {"input_15_7_pfm": ""},
            {"input_15_8_pfm": ""},
            {"input_15_9_pfm": ""},
            {"input_15_10_pfm": ""},
            {"input_15_11_pfm": ""},
            {"input_15_12_pfm": ""},
            {"input_15_13_pfm": ""},
            {"input_15_14_pfm": ""},
            {"input_15_15_pfm": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "ploplo.docx")
        output_file_path = os.path.join(
            script_directory, f"plo-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        a = DownloadModel(user_id=user_id, record_id=lastid, form_name="plo")
        a.save()

        end = time.time()
        print("End-Time: ", end)
        print("Time: ", end - start)

        response = FileResponse(open(output_file_path, "rb"))
        response["Content-Disposition"] = 'attachment; filename="plo-output.docx"'
        return response

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


@csrf_exempt
def plo_update(request):
    if request.method == "POST":
        # Process incoming data
        data = json.loads(request.body)
        user_id = data["user_id"]
        lastid = data["updateid"]

        inputvalues = [
            {"edit_values": "test values"},
            {"button_count": "test values"},
            {"auto_fetch_data": "test values"},
            {"prog_obj": ""},
            {"prog_desc": ""},
            {"ilo": ""},
            {"peo": ""},
            {"program": ""},
            {"college": ""},
            {"course-title-1": ""},
            {"course-title-2": ""},
            {"course-title-3": ""},
            {"course-title-4": ""},
            {"course-title-5": ""},
            {"course-title-6": ""},
            {"course-title-7": ""},
            {"course-title-8": ""},
            {"course-title-9": ""},
            {"course-title-10": ""},
            {"course-title-11": ""},
            {"course-title-12": ""},
            {"course-title-13": ""},
            {"course-title-14": ""},
            {"course-title-15": ""},
            {"plo-1": ""},
            {"plo-2": ""},
            {"plo-3": ""},
            {"plo-4": ""},
            {"plo-5": ""},
            {"plo-6": ""},
            {"plo-7": ""},
            {"plo-8": ""},
            {"plo-9": ""},
            {"plo-10": ""},
            {"plo-11": ""},
            {"plo-12": ""},
            {"plo-13": ""},
            {"plo-14": ""},
            {"plo-15": ""},
            {"assessment-tool-1": ""},
            {"assessment-tool-2": ""},
            {"assessment-tool-3": ""},
            {"assessment-tool-4": ""},
            {"assessment-tool-5": ""},
            {"assessment-tool-6": ""},
            {"assessment-tool-7": ""},
            {"assessment-tool-8": ""},
            {"assessment-tool-9": ""},
            {"assessment-tool-10": ""},
            {"assessment-tool-11": ""},
            {"assessment-tool-12": ""},
            {"assessment-tool-13": ""},
            {"assessment-tool-14": ""},
            {"assessment-tool-15": ""},
            {"tools-used-1": ""},
            {"tools-used-2": ""},
            {"tools-used-3": ""},
            {"tools-used-4": ""},
            {"tools-used-5": ""},
            {"tools-used-6": ""},
            {"tools-used-7": ""},
            {"tools-used-8": ""},
            {"tools-used-9": ""},
            {"tools-used-10": ""},
            {"tools-used-11": ""},
            {"tools-used-12": ""},
            {"tools-used-13": ""},
            {"tools-used-14": ""},
            {"tools-used-15": ""},
            {"students-1": ""},
            {"students-2": ""},
            {"students-3": ""},
            {"students-4": ""},
            {"students-5": ""},
            {"students-6": ""},
            {"students-7": ""},
            {"students-8": ""},
            {"students-9": ""},
            {"students-10": ""},
            {"students-11": ""},
            {"students-12": ""},
            {"students-13": ""},
            {"students-14": ""},
            {"students-15": ""},
            {"tool-passed-1": ""},
            {"tool-passed-2": ""},
            {"tool-passed-3": ""},
            {"tool-passed-4": ""},
            {"tool-passed-5": ""},
            {"tool-passed-6": ""},
            {"tool-passed-7": ""},
            {"tool-passed-8": ""},
            {"tool-passed-9": ""},
            {"tool-passed-10": ""},
            {"tool-passed-11": ""},
            {"tool-passed-12": ""},
            {"tool-passed-13": ""},
            {"tool-passed-14": ""},
            {"tool-passed-15": ""},
            {"course-passed-1": ""},
            {"course-passed-2": ""},
            {"course-passed-3": ""},
            {"course-passed-4": ""},
            {"course-passed-5": ""},
            {"course-passed-6": ""},
            {"course-passed-7": ""},
            {"course-passed-8": ""},
            {"course-passed-9": ""},
            {"course-passed-10": ""},
            {"course-passed-11": ""},
            {"course-passed-12": ""},
            {"course-passed-13": ""},
            {"course-passed-14": ""},
            {"course-passed-15": ""},
            {"input_1_0_pfm": ""},
            {"input_1_1_pfm": ""},
            {"input_1_2_pfm": ""},
            {"input_1_3_pfm": ""},
            {"input_1_4_pfm": ""},
            {"input_1_5_pfm": ""},
            {"input_1_6_pfm": ""},
            {"input_1_7_pfm": ""},
            {"input_1_8_pfm": ""},
            {"input_1_9_pfm": ""},
            {"input_1_10_pfm": ""},
            {"input_1_11_pfm": ""},
            {"input_1_12_pfm": ""},
            {"input_1_13_pfm": ""},
            {"input_1_14_pfm": ""},
            {"input_1_15_pfm": ""},
            {"input_2_0_pfm": ""},
            {"input_2_1_pfm": ""},
            {"input_2_2_pfm": ""},
            {"input_2_3_pfm": ""},
            {"input_2_4_pfm": ""},
            {"input_2_5_pfm": ""},
            {"input_2_6_pfm": ""},
            {"input_2_7_pfm": ""},
            {"input_2_8_pfm": ""},
            {"input_2_9_pfm": ""},
            {"input_2_10_pfm": ""},
            {"input_2_11_pfm": ""},
            {"input_2_12_pfm": ""},
            {"input_2_13_pfm": ""},
            {"input_2_14_pfm": ""},
            {"input_2_15_pfm": ""},
            {"input_3_0_pfm": ""},
            {"input_3_1_pfm": ""},
            {"input_3_2_pfm": ""},
            {"input_3_3_pfm": ""},
            {"input_3_4_pfm": ""},
            {"input_3_5_pfm": ""},
            {"input_3_6_pfm": ""},
            {"input_3_7_pfm": ""},
            {"input_3_8_pfm": ""},
            {"input_3_9_pfm": ""},
            {"input_3_10_pfm": ""},
            {"input_3_11_pfm": ""},
            {"input_3_12_pfm": ""},
            {"input_3_13_pfm": ""},
            {"input_3_14_pfm": ""},
            {"input_3_15_pfm": ""},
            {"input_4_0_pfm": ""},
            {"input_4_1_pfm": ""},
            {"input_4_2_pfm": ""},
            {"input_4_3_pfm": ""},
            {"input_4_4_pfm": ""},
            {"input_4_5_pfm": ""},
            {"input_4_6_pfm": ""},
            {"input_4_7_pfm": ""},
            {"input_4_8_pfm": ""},
            {"input_4_9_pfm": ""},
            {"input_4_10_pfm": ""},
            {"input_4_11_pfm": ""},
            {"input_4_12_pfm": ""},
            {"input_4_13_pfm": ""},
            {"input_4_14_pfm": ""},
            {"input_4_15_pfm": ""},
            {"input_5_0_pfm": ""},
            {"input_5_1_pfm": ""},
            {"input_5_2_pfm": ""},
            {"input_5_3_pfm": ""},
            {"input_5_4_pfm": ""},
            {"input_5_5_pfm": ""},
            {"input_5_6_pfm": ""},
            {"input_5_7_pfm": ""},
            {"input_5_8_pfm": ""},
            {"input_5_9_pfm": ""},
            {"input_5_10_pfm": ""},
            {"input_5_11_pfm": ""},
            {"input_5_12_pfm": ""},
            {"input_5_13_pfm": ""},
            {"input_5_14_pfm": ""},
            {"input_5_15_pfm": ""},
            {"input_6_0_pfm": ""},
            {"input_6_1_pfm": ""},
            {"input_6_2_pfm": ""},
            {"input_6_3_pfm": ""},
            {"input_6_4_pfm": ""},
            {"input_6_5_pfm": ""},
            {"input_6_6_pfm": ""},
            {"input_6_7_pfm": ""},
            {"input_6_8_pfm": ""},
            {"input_6_9_pfm": ""},
            {"input_6_10_pfm": ""},
            {"input_6_11_pfm": ""},
            {"input_6_12_pfm": ""},
            {"input_6_13_pfm": ""},
            {"input_6_14_pfm": ""},
            {"input_6_15_pfm": ""},
            {"input_7_0_pfm": ""},
            {"input_7_1_pfm": ""},
            {"input_7_2_pfm": ""},
            {"input_7_3_pfm": ""},
            {"input_7_4_pfm": ""},
            {"input_7_5_pfm": ""},
            {"input_7_6_pfm": ""},
            {"input_7_7_pfm": ""},
            {"input_7_8_pfm": ""},
            {"input_7_9_pfm": ""},
            {"input_7_10_pfm": ""},
            {"input_7_11_pfm": ""},
            {"input_7_12_pfm": ""},
            {"input_7_13_pfm": ""},
            {"input_7_14_pfm": ""},
            {"input_7_15_pfm": ""},
            {"input_8_0_pfm": ""},
            {"input_8_1_pfm": ""},
            {"input_8_2_pfm": ""},
            {"input_8_3_pfm": ""},
            {"input_8_4_pfm": ""},
            {"input_8_5_pfm": ""},
            {"input_8_6_pfm": ""},
            {"input_8_7_pfm": ""},
            {"input_8_8_pfm": ""},
            {"input_8_9_pfm": ""},
            {"input_8_10_pfm": ""},
            {"input_8_11_pfm": ""},
            {"input_8_12_pfm": ""},
            {"input_8_13_pfm": ""},
            {"input_8_14_pfm": ""},
            {"input_8_15_pfm": ""},
            {"input_9_0_pfm": ""},
            {"input_9_1_pfm": ""},
            {"input_9_2_pfm": ""},
            {"input_9_3_pfm": ""},
            {"input_9_4_pfm": ""},
            {"input_9_5_pfm": ""},
            {"input_9_6_pfm": ""},
            {"input_9_7_pfm": ""},
            {"input_9_8_pfm": ""},
            {"input_9_9_pfm": ""},
            {"input_9_10_pfm": ""},
            {"input_9_11_pfm": ""},
            {"input_9_12_pfm": ""},
            {"input_9_13_pfm": ""},
            {"input_9_14_pfm": ""},
            {"input_9_15_pfm": ""},
            {"input_10_0_pfm": ""},
            {"input_10_1_pfm": ""},
            {"input_10_2_pfm": ""},
            {"input_10_3_pfm": ""},
            {"input_10_4_pfm": ""},
            {"input_10_5_pfm": ""},
            {"input_10_6_pfm": ""},
            {"input_10_7_pfm": ""},
            {"input_10_8_pfm": ""},
            {"input_10_9_pfm": ""},
            {"input_10_10_pfm": ""},
            {"input_10_11_pfm": ""},
            {"input_10_12_pfm": ""},
            {"input_10_13_pfm": ""},
            {"input_10_14_pfm": ""},
            {"input_10_15_pfm": ""},
            {"input_11_0_pfm": ""},
            {"input_11_1_pfm": ""},
            {"input_11_2_pfm": ""},
            {"input_11_3_pfm": ""},
            {"input_11_4_pfm": ""},
            {"input_11_5_pfm": ""},
            {"input_11_6_pfm": ""},
            {"input_11_7_pfm": ""},
            {"input_11_8_pfm": ""},
            {"input_11_9_pfm": ""},
            {"input_11_10_pfm": ""},
            {"input_11_11_pfm": ""},
            {"input_11_12_pfm": ""},
            {"input_11_13_pfm": ""},
            {"input_11_14_pfm": ""},
            {"input_11_15_pfm": ""},
            {"input_12_0_pfm": ""},
            {"input_12_1_pfm": ""},
            {"input_12_2_pfm": ""},
            {"input_12_3_pfm": ""},
            {"input_12_4_pfm": ""},
            {"input_12_5_pfm": ""},
            {"input_12_6_pfm": ""},
            {"input_12_7_pfm": ""},
            {"input_12_8_pfm": ""},
            {"input_12_9_pfm": ""},
            {"input_12_10_pfm": ""},
            {"input_12_11_pfm": ""},
            {"input_12_12_pfm": ""},
            {"input_12_13_pfm": ""},
            {"input_12_14_pfm": ""},
            {"input_12_15_pfm": ""},
            {"input_13_0_pfm": ""},
            {"input_13_1_pfm": ""},
            {"input_13_2_pfm": ""},
            {"input_13_3_pfm": ""},
            {"input_13_4_pfm": ""},
            {"input_13_5_pfm": ""},
            {"input_13_6_pfm": ""},
            {"input_13_7_pfm": ""},
            {"input_13_8_pfm": ""},
            {"input_13_9_pfm": ""},
            {"input_13_10_pfm": ""},
            {"input_13_11_pfm": ""},
            {"input_13_12_pfm": ""},
            {"input_13_13_pfm": ""},
            {"input_13_14_pfm": ""},
            {"input_13_15_pfm": ""},
            {"input_14_0_pfm": ""},
            {"input_14_1_pfm": ""},
            {"input_14_2_pfm": ""},
            {"input_14_3_pfm": ""},
            {"input_14_4_pfm": ""},
            {"input_14_5_pfm": ""},
            {"input_14_6_pfm": ""},
            {"input_14_7_pfm": ""},
            {"input_14_8_pfm": ""},
            {"input_14_9_pfm": ""},
            {"input_14_10_pfm": ""},
            {"input_14_11_pfm": ""},
            {"input_14_12_pfm": ""},
            {"input_14_13_pfm": ""},
            {"input_14_14_pfm": ""},
            {"input_14_15_pfm": ""},
            {"input_15_0_pfm": ""},
            {"input_15_1_pfm": ""},
            {"input_15_2_pfm": ""},
            {"input_15_3_pfm": ""},
            {"input_15_4_pfm": ""},
            {"input_15_5_pfm": ""},
            {"input_15_6_pfm": ""},
            {"input_15_7_pfm": ""},
            {"input_15_8_pfm": ""},
            {"input_15_9_pfm": ""},
            {"input_15_10_pfm": ""},
            {"input_15_11_pfm": ""},
            {"input_15_12_pfm": ""},
            {"input_15_13_pfm": ""},
            {"input_15_14_pfm": ""},
            {"input_15_15_pfm": ""},
            {"id": ""},
            {"form_id": ""},
        ]

        input_w_values_data = data["input_w_values_data"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "plo.docx")
        output_file_path = os.path.join(
            script_directory, f"plo-{user_id}-{lastid}.docx"
        )
        doc = Document(input_file_path)

        replacement_dict = {
            key: value for item in inputvalues for key, value in item.items()
        }
        for data_item in input_w_values_data:
            replacement_dict.update(data_item)

        replace_text_v1(doc, replacement_dict)

        save_document_with_temp(doc, output_file_path)

        try:
            DownloadModel.objects.update_or_create(
                user_id=user_id, record_id=lastid, form_name="plo", defaults={}
            )
        except DownloadModel.MultipleObjectsReturned:
            DownloadModel.objects.filter(
                user_id=user_id, record_id=lastid, form_name="plo"
            ).delete()
            DownloadModel.objects.get_or_create(
                user_id=user_id, record_id=lastid, form_name="plo"
            )

        return JsonResponse({"status": "success", "message": "File created"})

    else:
        return JsonResponse({"status": "error", "message": "Invalid request method"})


import os
import shutil
import openpyxl
from openpyxl import Workbook
from django.http import HttpResponse
from openpyxl.styles import Font
from decimal import Decimal


@csrf_exempt
def class_record(request):
    if request.method == "POST":
        data = json.loads(request.body)

        action = data["action"]
        user_id = data["user_id"]
        lastid = 0

        if action == "save":
            lastid = data["lastid"]
        elif action == "update":
            lastid = data["updateid"]
            DownloadModel.objects.filter(
                user_id=user_id, record_id=lastid, form_name="class-record"
            ).delete()

        input_w_values_data = data["input_w_values_data"]
        student_count = data["student_count"]
        csprelim = data["csprelim"]
        csmidterm = data["csmidterm"]
        csfinal = data["csfinal"]
        prelim_exam = data["prelim_exam"]
        midterm_exam = data["midterm_exam"]
        final_exam = data["final_exam"]

        script_directory = os.path.dirname(os.path.abspath(__file__))
        input_file_path = os.path.join(script_directory, "Proposed-Class-record.xlsx")

        # Define paths
        duplicate_file_path = os.path.join(
            script_directory, f"Class-record-{user_id}-{lastid}.xlsx"
        )
        if os.path.exists(duplicate_file_path):
            os.remove(duplicate_file_path)
        shutil.copy2(input_file_path, duplicate_file_path)

        try:
            workbook = openpyxl.load_workbook(duplicate_file_path)
        except FileNotFoundError:
            workbook = Workbook()

        # Select active worksheet or create a new one
        worksheet = workbook.active

        student_count = int(student_count) + 1

        letters = [
            "",
            "F",
            "G",
            "H",
            "I",
            "J",
            "K",
            "L",
            "M",
            "N",
            "O",
            "P",
            "Q",
            "R",
            "S",
            "T",
            "U",
            "V",
            "W",
            "X",
            "Y",
            "Z",
            "AA",
            "AB",
            "AC",
            "AD",
            "AE",
            "AF",
            "AG",
            "AH",
            "AI",
            "AJ",
            "AK",
            "AL",
            "AM",
            "AN",
            "AO",
        ]
        letters2 = [
            "",
            "Q",
            "R",
            "S",
            "T",
            "U",
            "V",
            "W",
            "X",
            "Y",
            "Z",
            "AA",
            "AB",
            "AC",
            "AD",
            "AE",
            "AF",
            "AG",
            "AH",
            "AI",
            "AJ",
            "AK",
            "AL",
            "AM",
            "AN",
            "AO",
        ]
        letters3 = [
            "",
            "AC",
            "AD",
            "AE",
            "AF",
            "AG",
            "AH",
            "AI",
            "AJ",
            "AK",
            "AL",
            "AM",
            "AN",
            "AO",
            "AP",
        ]

        # BASIC INFORMATION HEADER =================
        try:
            cell_address2 = f"E5"
            key_to_search = f"academic_year"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    break
            cell_address2 = f"E6"
            key_to_search = f"schedule"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    break
            cell_address2 = f"C5"
            key_to_search = f"sections"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    break
            cell_address2 = f"C6"
            key_to_search = f"course_code"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    break
            cell_address2 = f"C7"
            key_to_search = f"course_title"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    break

            # PRELIMS HEADER =================
            prelim_headers = [""]
            cell_address2 = f"F11"
            key_to_search = f"prelim-total-SW-1"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    prelim_headers.append(item[key_to_search])
                    break
            cell_address2 = f"G11"
            key_to_search = f"prelim-total-SW-2"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    prelim_headers.append(item[key_to_search])
                    break
            cell_address2 = f"H11"
            key_to_search = f"prelim-total-SW-3"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    prelim_headers.append(item[key_to_search])
                    break

            cell_address2 = f"I11"
            prelim_sum_quiz = sum(int(x) for x in prelim_headers if x.isdigit())
            worksheet[cell_address2] = prelim_sum_quiz

            cell_address2 = f"J10"
            worksheet[cell_address2] = csprelim + "%"
            cell_address2 = f"K11"
            worksheet[cell_address2] = prelim_exam
            re_less_100 = 100 - int(csprelim)
            cell_address2 = f"L10"
            worksheet[cell_address2] = str(re_less_100) + "%"

            # MIDTERM HEADER =================
            midterm_headers = [""]
            cell_address2 = f"Q11"
            key_to_search = f"midterm-total-SW-1"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    midterm_headers.append(item[key_to_search])
                    break
            cell_address2 = f"R11"
            key_to_search = f"midterm-total-SW-2"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    midterm_headers.append(item[key_to_search])
                    break
            cell_address2 = f"S11"
            key_to_search = f"midterm-total-SW-3"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    midterm_headers.append(item[key_to_search])
                    break
            cell_address2 = f"T11"
            midterm_sum_quiz = sum(int(x) for x in midterm_headers if x.isdigit())
            worksheet[cell_address2] = midterm_sum_quiz

            cell_address2 = f"U10"
            key_to_search = f"csmidterm"
            less_csmidterm = 0
            for item in input_w_values_data:
                if key_to_search in item:
                    lmidterm = 100 - int(item[key_to_search])
                    less_csmidterm = lmidterm
                    worksheet[cell_address2] = item[key_to_search] + "%"
                    break
            cell_address2 = f"V11"
            key_to_search = f"midterm-exam"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    break
            cell_address2 = f"W10"
            worksheet[cell_address2] = str(less_csmidterm) + "%"

            # FINALS HEADER =================
            finals_headers = [""]
            cell_address2 = f"AC11"
            key_to_search = f"final-total-SW-1"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    finals_headers.append(item[key_to_search])
                    break
            cell_address2 = f"AD11"
            key_to_search = f"final-total-SW-2"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    finals_headers.append(item[key_to_search])
                    break
            cell_address2 = f"AE11"
            key_to_search = f"final-total-SW-3"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search]
                    finals_headers.append(item[key_to_search])
                    break
            cell_address2 = f"AF11"
            final_sum_quiz = sum(int(x) for x in finals_headers if x.isdigit())
            worksheet[cell_address2] = final_sum_quiz

            final_less_100 = 0
            cell_address2 = f"AG10"
            key_to_search = f"csfinal"
            for item in input_w_values_data:
                if key_to_search in item:
                    worksheet[cell_address2] = item[key_to_search] + "%"
                    final_less_100 = int(100) - int(item[key_to_search])
                    break

            cell_address2 = f"AI10"
            worksheet[cell_address2] = str(final_less_100) + "%"

            cell_address2 = f"AH11"
            worksheet[cell_address2] = final_exam

        except:
            pass

        quiz_list = []
        row = 12
        student_exam = 0
        quizes_results = 0
        pe_exam_result = 0
        prelim_grade = 0
        student_quiz_got_prelim = []
        student_quiz_got_midterm = []
        student_quiz_got_finals = []
        final_prelim_grade = 0
        final_midterm_grade = 0
        final_final_grade = 0

        # STUDENTS PRELIMS INFORMATION =================
        for i in range(1, student_count):
            student_quiz_got_prelim = []
            for j in range(1, 12):
                cell_address2 = f"{letters[j]}{row}"
                key_to_search = f"prelim-student-SW-{j}-{i}"
                for item in input_w_values_data:
                    if key_to_search in item:
                        try:
                            worksheet[cell_address2] = (
                                item[key_to_search] + "/" + prelim_headers[j]
                            )
                            quiz_list.append(item[key_to_search])
                            student_quiz_got_prelim.append(item[key_to_search])
                        except:
                            pass
                        break

                if j == 4:
                    sum_of_quiz = sum(int(x) for x in prelim_headers if x.isdigit())
                    sum_of_quiz_student_got = sum(
                        int(x) for x in student_quiz_got_prelim if x.isdigit()
                    )
                    worksheet[cell_address2] = (
                        str(sum_of_quiz_student_got) + "/" + str(sum_of_quiz)
                    )

                if j == 5:
                    sum_of_quiz = sum(
                        int(x) for x in student_quiz_got_prelim if x.isdigit()
                    )
                    sum_of_quiz_headers = sum(
                        int(x) for x in prelim_headers if x.isdigit()
                    )
                    division_res = (int(sum_of_quiz) / int(sum_of_quiz_headers)) * int(
                        csprelim
                    )
                    division_res = round(division_res, 2)
                    quizes_results = division_res
                    worksheet[cell_address2] = division_res

                if j == 6:
                    key_to_search = f"prelim-exam-student-{i}"
                    for item in input_w_values_data:
                        if key_to_search in item:
                            worksheet[cell_address2] = (
                                item[key_to_search] + "/" + prelim_exam
                            )
                            student_exam = item[key_to_search]
                            break
                if j == 7:
                    re_less_100 = int(100) - int(csprelim)
                    division_res = (int(student_exam) / int(prelim_exam)) * int(
                        re_less_100
                    )
                    pe_exam_result = division_res
                    worksheet[cell_address2] = division_res
                    quiz_list = []

                if j == 8:
                    key_to_search = f"clo1-student-{i}"
                    for item in input_w_values_data:
                        if key_to_search in item:
                            worksheet[cell_address2] = item[key_to_search]
                            student_exam = item[key_to_search]
                            break

                if j == 9:
                    key_to_search = f"clo1-attained-student-{i}"
                    for item in input_w_values_data:
                        if key_to_search in item:
                            worksheet[cell_address2] = item[key_to_search]
                            student_exam = item[key_to_search]
                            break

                if j == 10:
                    prelim_grade = Decimal(quizes_results) + Decimal(pe_exam_result)
                    worksheet[cell_address2] = prelim_grade
                    final_prelim_grade = prelim_grade

                if j == 11:
                    if prelim_grade >= 75:
                        worksheet[cell_address2] = "Passed"
                    else:
                        worksheet[cell_address2] = "Failed"

            # STUDENTS MIDTERM INFORMATION =================
            for j in range(1, 13):
                cell_address2 = f"{letters2[j]}{row}"
                key_to_search = f"midterm-student-SW-{j}-{i}"
                for item in input_w_values_data:
                    if key_to_search in item:
                        try:
                            worksheet[cell_address2] = (
                                item[key_to_search] + "/" + midterm_headers[j]
                            )
                            quiz_list.append(item[key_to_search])
                            student_quiz_got_midterm.append(item[key_to_search])
                        except:
                            pass
                        break

                if j == 4:
                    sum_of_quiz = sum(int(x) for x in midterm_headers if x.isdigit())
                    sum_of_quiz_student_got = sum(
                        int(x) for x in student_quiz_got_midterm if x.isdigit()
                    )
                    worksheet[cell_address2] = (
                        str(sum_of_quiz_student_got) + "/" + str(sum_of_quiz)
                    )

                if j == 5:
                    sum_of_quiz = sum(
                        int(x) for x in student_quiz_got_midterm if x.isdigit()
                    )
                    sum_of_quiz_headers = sum(
                        int(x) for x in midterm_headers if x.isdigit()
                    )
                    division_res = (int(sum_of_quiz) / int(sum_of_quiz_headers)) * int(
                        csmidterm
                    )
                    division_res = round(division_res, 2)
                    quizes_results = division_res
                    worksheet[cell_address2] = division_res

                if j == 6:
                    key_to_search = f"mid-exam-student-{i}"
                    for item in input_w_values_data:
                        if key_to_search in item:
                            worksheet[cell_address2] = (
                                item[key_to_search] + "/" + midterm_exam
                            )
                            student_exam = item[key_to_search]
                            break

                if j == 7:
                    re_less_100 = int(100) - int(csmidterm)
                    division_res = (int(student_exam) / int(midterm_exam)) * int(
                        re_less_100
                    )
                    pe_exam_result = division_res
                    worksheet[cell_address2] = division_res
                    quiz_list = []

                if j == 8:
                    key_to_search = f"clo2-student-{i}"
                    for item in input_w_values_data:
                        if key_to_search in item:
                            worksheet[cell_address2] = item[key_to_search]
                            break

                if j == 9:
                    key_to_search = f"clo2-attained-student-{i}"
                    for item in input_w_values_data:
                        if key_to_search in item:
                            worksheet[cell_address2] = item[key_to_search]
                            break

                if j == 10:
                    midterm_grade = Decimal(quizes_results) + Decimal(pe_exam_result)
                    worksheet[cell_address2] = midterm_grade
                    final_midterm_grade = midterm_grade

                if j == 11:
                    cqi = Decimal(final_prelim_grade) + Decimal(final_midterm_grade)
                    cqi /= 2  # Divide by 2
                    worksheet[cell_address2] = cqi

                if j == 12:
                    cqi = Decimal(final_prelim_grade) + Decimal(final_midterm_grade)
                    cqi /= 2  # Divide by 2
                    if cqi >= 80:
                        worksheet[cell_address2] = "No"
                    else:
                        worksheet[cell_address2] = "Yes"

            # STUDENTS FINALS INFORMATION =================
            for j in range(1, 15):
                cell_address2 = f"{letters3[j]}{row}"
                key_to_search = f"final-student-SW-{j}-{i}"
                for item in input_w_values_data:
                    if key_to_search in item:
                        try:
                            worksheet[cell_address2] = (
                                item[key_to_search] + "/" + finals_headers[j]
                            )
                            quiz_list.append(item[key_to_search])
                            student_quiz_got_finals.append(item[key_to_search])
                        except:
                            pass
                        break

                if j == 4:
                    sum_of_quiz = sum(int(x) for x in finals_headers if x.isdigit())
                    sum_of_quiz_student_got = sum(
                        int(x) for x in student_quiz_got_finals if x.isdigit()
                    )
                    worksheet[cell_address2] = (
                        str(sum_of_quiz_student_got) + "/" + str(sum_of_quiz)
                    )

                if j == 5:
                    sum_of_quiz = sum(
                        int(x) for x in student_quiz_got_finals if x.isdigit()
                    )
                    sum_of_quiz_headers = sum(
                        int(x) for x in finals_headers if x.isdigit()
                    )
                    division_res = (int(sum_of_quiz) / int(sum_of_quiz_headers)) * int(
                        csfinal
                    )
                    division_res = round(division_res, 2)
                    quizes_results = division_res
                    worksheet[cell_address2] = division_res

                if j == 6:
                    key_to_search = f"final-exam-student-{i}"
                    for item in input_w_values_data:
                        if key_to_search in item:
                            worksheet[cell_address2] = (
                                item[key_to_search] + "/" + final_exam
                            )
                            student_exam = item[key_to_search]
                            break

                if j == 7:
                    re_less_100 = int(100) - int(csfinal)
                    division_res = (int(student_exam) / int(final_exam)) * int(
                        re_less_100
                    )
                    pe_exam_result = division_res
                    worksheet[cell_address2] = division_res

                if j == 8:
                    key_to_search = f"clo3-student-{i}"
                    for item in input_w_values_data:
                        if key_to_search in item:
                            worksheet[cell_address2] = item[key_to_search]
                            break

                if j == 9:
                    key_to_search = f"clo3-attained-student-{i}"
                    for item in input_w_values_data:
                        if key_to_search in item:
                            worksheet[cell_address2] = item[key_to_search]
                            break

                if j == 10:
                    final_final_grade = Decimal(quizes_results) + Decimal(
                        pe_exam_result
                    )
                    worksheet[cell_address2] = final_final_grade

                if j == 11:
                    sem_grade = (
                        (
                            Decimal(final_prelim_grade)
                            + Decimal(final_midterm_grade)
                            + Decimal(final_final_grade)
                        )
                        / 3
                    ) + 5
                    sem_grade = round(sem_grade, 2)
                    worksheet[cell_address2] = sem_grade

                if j == 12:
                    sem_grade = (
                        (
                            Decimal(final_prelim_grade)
                            + Decimal(final_midterm_grade)
                            + Decimal(final_final_grade)
                        )
                        / 3
                    ) + 5
                    sem_grade = round(sem_grade, 2)

                    if sem_grade >= 96.96 and sem_grade <= 100:
                        grade = "1.00"
                    elif sem_grade >= 93.35 and sem_grade <= 96.68:
                        grade = "1.25"
                    elif sem_grade >= 90.01 and sem_grade <= 93.34:
                        grade = "1.50"
                    elif sem_grade >= 86.68 and sem_grade <= 90.00:
                        grade = "1.75"
                    elif sem_grade >= 83.35 and sem_grade <= 86.67:
                        grade = "2.00"
                    elif sem_grade >= 80.01 and sem_grade <= 83.34:
                        grade = "2.25"
                    elif sem_grade >= 76.68 and sem_grade <= 80.00:
                        grade = "2.50"
                    elif sem_grade >= 73.34 and sem_grade <= 76.67:
                        grade = "2.75"
                    elif sem_grade >= 70.00 and sem_grade <= 73.33:
                        grade = "3.00"
                    elif sem_grade < 70:
                        grade = "5.00"
                    else:
                        grade = "Dropped"

                    worksheet[cell_address2] = grade

                if j == 13:
                    sem_grade = (
                        (
                            Decimal(final_prelim_grade)
                            + Decimal(final_midterm_grade)
                            + Decimal(final_final_grade)
                        )
                        / 3
                    ) + 5
                    sem_grade = round(sem_grade, 2)

                    if sem_grade >= 70:
                        final_remarks = "Passed"
                    elif sem_grade < 70:
                        final_remarks = "Failed"
                    else:
                        final_remarks = "Dropped"

                    worksheet[cell_address2] = final_remarks

            student_quiz_got_midterm = []
            student_quiz_got_finals = []
            row += 1

        target_row = 12
        for item in input_w_values_data:
            for key, value in item.items():
                if "student-sn" in key:
                    cell_address = f"B{target_row}"
                    font = Font(size=10, bold=False)
                    worksheet[cell_address] = str(value)
                    worksheet[cell_address].font = font
                    target_row += 1
                elif "student-ln" in key:
                    cell_address = f"C{target_row}"
                    worksheet[cell_address] = value
                elif "student-fn" in key:
                    cell_address = f"D{target_row}"
                    worksheet[cell_address] = value
                elif "student-mn" in key:
                    cell_address = f"E{target_row}"
                    va = value[0]
                    worksheet[cell_address] = va

        workbook.save(duplicate_file_path)

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = (
            f"attachment; filename={os.path.basename(duplicate_file_path)}"
        )
        workbook.save(response)

        a = DownloadModel(user_id=user_id, record_id=lastid, form_name="class-record")
        a.save()

        return response
    return HttpResponse("Nothing to download")


def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if old_text in run.text:
                new_run = paragraph.add_run()
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                new_run.font.color.rgb = run.font.color.rgb
                new_run.text = new_text
                run.clear()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            new_run = paragraph.add_run()
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.size = run.font.size
                            new_run.font.name = run.font.name
                            new_run.font.color.rgb = run.font.color.rgb
                            new_run.text = new_text
                            run.clear()
